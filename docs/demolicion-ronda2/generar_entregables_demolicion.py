# -*- coding: utf-8 -*-
"""Extrae datos de informes demolición, cruza Habitable, genera Excel + Word."""
from __future__ import annotations

import re
import unicodedata
from datetime import date
from pathlib import Path

import fitz
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

SRC_PDF = Path(r"C:\Users\PC\Downloads\Edificios Demolicion")
CSV_HAB = Path(r"C:\Users\PC\Downloads\habitable_inspecciones_2026-08-20_12-02-48.csv")
OUT = Path(__file__).resolve().parent
EXTRACT = OUT / "_extract"
EXTRACT.mkdir(parents=True, exist_ok=True)

# Preferencias de cruce por patrón de archivo → id Habitable (revisión asistida).
# Vacío = solo heurística. Varios candidatos → dejar que la heurística elija y marcar revisión.
MANUAL_HAB_ID: dict[str, int | None] = {
    "samantha": 195050,
    "arenal": 77330,
    "el mamon": 170487,
    "refugio del caribe": 170205,
    "franco mar": 171818,
    "francomar": 171818,
    "camuri beach": 95703,
    "contry mar": 162009,
    "country mar": 162009,
    "dictis": 157349,
    "ibiza": 159293,
    "medusa": 70266,
    "parque azul": 161871,
    "sol de oro ii": 140599,
    "sol de oro iv": 190094,
    "tanaguarena mar": 170020,
    "castromar": 32666,
    "arrecife": 68099,
    "brisamar": 131643,
    "altomar": 114439,
}


def strip_accents(s: str) -> str:
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return ""
    s = str(s)
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c)).lower()


def norm_name(s: str) -> str:
    s = strip_accents(s)
    s = re.sub(r"[^a-z0-9\s]", " ", s)
    for w in (
        "informe",
        "tecnico",
        "evaluacion",
        "residencias",
        "residencia",
        "edificio",
        "edificios",
        "justificacion",
        "de",
        "demolicion",
        "inspeccion",
        "final",
        "firmado",
        "la",
        "el",
        "los",
        "las",
        "del",
        "pdf",
        "hotel",
    ):
        s = re.sub(rf"\b{w}\b", " ", s)
    return re.sub(r"\s+", " ", s).strip()


def extract_pdf_text(path: Path, max_pages: int = 12) -> tuple[str, int]:
    doc = fitz.open(path)
    n = doc.page_count
    parts = [doc[i].get_text("text") for i in range(min(n, max_pages))]
    if n > max_pages:
        parts.append(doc[n - 1].get_text("text"))
    text = "\n".join(parts)
    text = text.replace("\u200b", "").replace("\ufeff", "").replace("\u00a0", " ")
    return text, n


def first_match(patterns, text, flags=re.I | re.S):
    for p in patterns:
        m = re.search(p, text, flags)
        if m:
            g = m.group(1) if m.lastindex else m.group(0)
            return re.sub(r"\s+", " ", g).strip()
    return ""


def extract_fields(fname: str, text: str, pages: int) -> dict:
    t = text
    lower = t.lower()

    stem = Path(fname).stem
    nombre_archivo = re.sub(
        r"(?i)^(informe\s*(tecnico|evaluacion)?|justificaci[oó]n\s*de\s*demolici[oó]n|"
        r"inspeccion\s*-?\s*|informe\s*demolicion)\s*",
        "",
        stem,
    ).strip(" -_")
    nombre_archivo = re.sub(r"\s*\(\d+\)\s*$", "", nombre_archivo)
    nombre_archivo = re.sub(r"\s+\d{10,}.*$", "", nombre_archivo)

    nombre = first_match(
        [
            r"Edificaci[oó]n\s*Inspeccionada\s*:\s*([^\n]+)",
            r"Nombre de la edificaci[oó]n[^\n]*:\s*\n?\s*([^\n]+)",
            r"Edificaci[oó]n\s*\n([^\n]+)",
            r"Proyecto\s*:\s*([^\n]+)",
            r"Edificio\s+([A-ZÁÉÍÓÚÑ0-9][^\n]{3,60})",
            r"denominada\s+[\"“]?([^\"”\n]{4,80})[\"”]?",
        ],
        t,
    )
    if not nombre or len(nombre) < 4:
        nombre = nombre_archivo

    ubicacion = first_match(
        [
            r"Ubicaci[oó]n\s*:\s*([^\n]+)",
            r"Ubicaci[oó]n\s*\n([^\n]+)",
        ],
        t,
    )

    lat = lon = ""
    m = re.search(
        r"(?:Latitud|Lat)\s*[:=]?\s*(?:N\s*)?([+-]?\d{1,2}[.,]\d{3,8}).{0,40}?"
        r"(?:Longitud|Lon|Lng)\s*[:=]?\s*(?:W\s*)?(-?\d{1,2}[.,]\d{3,8})",
        t,
        re.I | re.S,
    )
    if m:
        lat, lon = m.group(1).replace(",", "."), m.group(2).replace(",", ".")
        if not lon.startswith("-") and "W" in t[m.start() : m.end()].upper():
            lon = "-" + lon.lstrip("-")
    else:
        m2 = re.search(
            r"(N\s*)?(\d{1,2}[.,]\d{4,8})\s*[,°]?\s*(?:W|Oeste)?\s*(-?\d{1,2}[.,]\d{4,8})",
            t,
            re.I,
        )
        if m2:
            lat = m2.group(2).replace(",", ".")
            lon = m2.group(3).replace(",", ".")
            if not lon.startswith("-"):
                lon = "-" + lon

    doc_id = first_match(
        [
            r"(DGPS[^\s\n/]{0,5}[-_/]?\d{0,4}[-_/]?\d{0,2}[-_/]?[^\s\n]{0,25})",
            r"Documento\s*N[°º\.]?\s*([^\n]+)",
        ],
        t,
    )

    fecha = first_match(
        [
            r"(?:Fecha(?:\s+de\s+(?:la\s+)?(?:evaluaci[oó]n|inspecci[oó]n))?|Evaluaci[oó]n)\s*[:=]?\s*"
            r"(\[?\s*\d{1,2}\s*/\s*\d{1,2}\s*/\s*2026\s*\]?|\d{1,2}\s+de\s+\w+\s+de\s+2026)",
            r"(\d{1,2}\s+de\s+(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|"
            r"octubre|noviembre|diciembre)\s+(?:de\s+)?2026)",
            r"(\d{1,2}[/-]\d{1,2}[/-]2026)",
            r"(La Guaira,?\s+\d{1,2}\s+de\s+\w+\s+2026)",
        ],
        t,
    )

    evals = []
    for m in re.finditer(
        r"Ing\.?\s+([A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÑáéíóúñ\.\s']{3,45})"
        r"(?:\s*\(?\s*C\.?I\.?\s*[\d\.]+)?",
        t,
    ):
        name = re.sub(r"\s+", " ", m.group(1)).strip(" .")
        if len(name) >= 5:
            evals.append(name)
    seen: set[str] = set()
    evaluadores = []
    for e in evals:
        k = strip_accents(e)
        if k not in seen:
            seen.add(k)
            evaluadores.append(e)

    supervisor = first_match(
        [
            r"Supervisor(?:\s+Institucional)?\s*[:=]?\s*(?:Ing\.?\s*)?([^\n]+)",
            r"Ingeniero Supervisor[^\n]*:\s*(?:Ing\.?\s*)?([^\n]+)",
        ],
        t,
    )

    if "pnud" in lower or "atc-20" in lower:
        tipo = "Ficha PNUD / ATC-20 detallada"
    elif "recolecci" in lower and "escombro" in lower:
        tipo = "Recolección de escombros / post-colapso"
    elif "justificaci" in lower and "demolic" in lower:
        tipo = "Justificación técnica de demolición"
    elif "nivel 2" in lower or ("detallada" in lower and "posts" in lower):
        tipo = "Informe técnico nivel 2 (detallado)"
    else:
        tipo = "Informe técnico estructural"

    if re.search(r"tarjeta\s+roja|etiqueta\s+roja|inseguro\s*\(tarjeta", lower):
        dictamen = "INSEGURO (Rojo) — confirmado/citado"
    elif "colaps" in lower and "escombro" in lower:
        dictamen = "Colapso / escombros"
    elif "demolic" in lower:
        dictamen = "Dictamen orientado a demolición (sin etiqueta explícita en texto)"
    else:
        dictamen = ""

    recomienda_dem = bool(
        re.search(
            r"demolici[oó]n\s+(r[aá]pida|controlada|total|t[eé]cnica)|"
            r"proceder\s+con\s+(su\s+)?demolici|"
            r"[uú]nica\s+acci[oó]n\s+de\s+mitigaci[oó]n|"
            r"tramitar\s+la\s+orden\s+formal\s+de\s+demolici|"
            r"necesidad\s+inminente\s+de\s+proceder\s+con\s+la\s+demolici",
            lower,
        )
    )
    inviabilidad_rep = bool(
        re.search(
            r"inviabilidad\s+de\s+repar|no\s+existe\s+viabilidad|"
            r"irrecuperable|invalidan\s+t[eé]cnica",
            lower,
        )
    )

    n_pisos = first_match(
        [
            r"(\d{1,2})\s+[Nn]iveles",
            r"(\d{1,2})\s+pisos\s+de\s+altura",
            r"Caracter[ií]sticas\s+geom[eé]tricas\s*:\s*(\d+)\s+Niveles",
        ],
        t,
    )
    word_map = {
        "catorce": "14",
        "trece": "13",
        "doce": "12",
        "once": "11",
        "diez": "10",
        "nueve": "9",
        "ocho": "8",
        "siete": "7",
        "seis": "6",
        "cinco": "5",
    }
    if not n_pisos:
        m = re.search(
            r"(catorce|trece|doce|once|diez|nueve|ocho|siete|seis|cinco)\s*\(?(\d{0,2})\)?\s*pisos",
            lower,
        )
        if m:
            n_pisos = m.group(2) or word_map.get(m.group(1), "")

    n_sotanos = first_match(
        [
            r"(\d+)\s+niveles?\s+de\s+s[oó]tano",
            r"(\d+)\s+S[oó]tanos",
            r"y\s+(\d+)\s+niveles?\s+de\s+s[oó]tano",
        ],
        t,
    )
    area_m2 = first_match(
        [
            r"[Aa]rea(?:\s+de\s+planta|\s+[Tt]otal)?[^\d]{0,30}(\d{2,4}(?:[.,]\d+)?)\s*m",
            r"(\d{2,4}(?:[.,]\d+)?)\s*m[²2]",
        ],
        t,
    )
    sistema = first_match(
        [
            r"[Ss]istema\s+estructural[^:\n]*:\s*([^\n.]+)",
            r"(p[oó]rticos?\s+de\s+concreto\s+armado[^\n.]{0,40})",
        ],
        t,
    )
    uso = first_match(
        [
            r"[Uu]so(?:\s+ocupacional)?(?:\s+principal)?\s*[:=]\s*([^\n]+)",
            r"[Uu]so\s+y\s+[Cc]onfiguraci[oó]n\s*:\s*([^\n.]+)",
        ],
        t,
    )
    anio = first_match(
        [
            r"[Aa][nñ]o\s+de\s+construcci[oó]n\s*[:=]\s*([^\n]+)",
            r"(19\d{2}|20\d{2})\s*[-–]\s*(19\d{2}|20\d{2})",
        ],
        t,
    )
    piso_critico = first_match(
        [
            r"[Pp]iso\s+[Cc]r[ií]tico[^\n]{0,40}",
            r"[Ee]ntrepiso\s+[Cc]r[ií]tico[^\n]{0,60}",
            r"(Nivel\s+1|Planta\s+Baja\s*/\s*Nivel\s+PB)[^\n]{0,40}",
        ],
        t,
    )

    flags = []
    checks = [
        ("falla_columnas", r"falla.*(columna|compresi)|aplastamiento|estallido.*(n[uú]cleo|concreto)|pandeo"),
        ("perdida_verticalidad", r"p[eé]rdida\s+de\s+verticalidad|inclinaci[oó]n|desplome|p-delta"),
        ("dano_vigas", r"grieta.*(viga|corte|45)|deformaci[oó]n.*viga"),
        ("dano_losas", r"losa|volado|horizontalidad"),
        ("cerramientos", r"mamposter[ií]a|cerramiento|fachada"),
        ("peligro_aledanos", r"aleda[nñ]|vecin|colindan|15\s*a\s*20\s*metros"),
        ("acordonamiento", r"acordon|exclusi[oó]n|prohibici[oó]n\s+de\s+ingreso|no\s+entre"),
        ("apuntalamiento", r"apuntal"),
        ("monitoreo_vecinos", r"monitoreo|instrumentaci[oó]n\s+geod"),
    ]
    for key, pat in checks:
        if re.search(pat, lower):
            flags.append(key)

    concl = ""
    m = re.search(r"(?:4\.?\s*1\.?\s*)?[Cc]onclusiones[^\n]*\n(.{80,600})", t, re.S)
    if m:
        concl = re.sub(r"\s+", " ", m.group(1))[:500]
    elif recomienda_dem:
        m2 = re.search(r"([^\n]{0,120}demolici[oó]n[^\n]{0,200})", t, re.I)
        if m2:
            concl = re.sub(r"\s+", " ", m2.group(1))[:500]

    fotos = len(re.findall(r"[Ff]otograf[ií]a\s*\d|[Ee]videncia\s+[Ff]otogr[aá]fica\s*N", t))

    if recomienda_dem:
        rec_dem = "Sí"
    elif "demolic" in lower:
        rec_dem = "Parcial/implícito"
    else:
        rec_dem = "No detectado en texto"

    return {
        "archivo_pdf": fname,
        "paginas_pdf": pages,
        "tipo_informe": tipo,
        "codigo_documento": doc_id,
        "nombre_edificacion_informe": nombre,
        "nombre_norm": norm_name(nombre),
        "nombre_archivo_norm": norm_name(nombre_archivo),
        "ubicacion_informe": (ubicacion or "")[:200],
        "lat_informe": lat,
        "lon_informe": lon,
        "fecha_inspeccion_informe": fecha,
        "evaluadores": "; ".join(evaluadores[:8]),
        "supervisor": (supervisor or "")[:120],
        "dictamen_etiqueta": dictamen,
        "recomienda_demolicion": rec_dem,
        "inviabilidad_reparacion": "Sí" if inviabilidad_rep else "No detectado",
        "uso": (uso or "")[:120],
        "anio_construccion_informe": (anio or "")[:40],
        "num_pisos_informe": n_pisos,
        "num_sotanos_informe": n_sotanos,
        "area_planta_m2": area_m2,
        "sistema_estructural": (sistema or "")[:160],
        "piso_critico_mencionado": (piso_critico or "")[:120],
        "hallazgos_flags": "; ".join(flags),
        "n_evidencias_fotograficas_citadas": fotos,
        "extracto_conclusiones": concl,
        "chars_texto_extraido": len(t),
        "texto_disponible": "Sí" if len(t) > 200 else "Escaso (posible PDF escaneado)",
    }


def fuzzy_tokens(a: str, b: str) -> float:
    ta = set(a.split())
    tb = set(b.split())
    if not ta or not tb:
        return 0.0
    return len(ta & tb) / max(len(ta), len(tb))


def _score_row(row: dict, hn: str, hd: str, lat, lng, keys, tokens) -> float:
    score = 0.0
    for k in keys:
        score = max(score, fuzzy_tokens(k, hn))
        if k and hn and (k in hn or hn in k):
            score = max(score, 0.9)
    if tokens:
        hit = sum(1 for tok in tokens if tok in hn or tok in hd)
        score = max(score, hit / max(3, len(tokens)) * 0.7 + (0.2 if hit >= 2 else 0))
    if row["lat_informe"] and row["lon_informe"] and pd.notna(lat) and pd.notna(lng):
        try:
            dlat = abs(float(row["lat_informe"]) - float(lat))
            dlon = abs(float(row["lon_informe"]) - float(lng))
            if dlat < 0.0015 and dlon < 0.0015:
                score = max(score, 0.95)
            elif dlat < 0.004 and dlon < 0.004:
                score = max(score, 0.75)
        except Exception:
            pass
    return score


def best_hab_matches(row: dict, hab_pool: pd.DataFrame, top_n: int = 3) -> list[dict]:
    """Match against a prefiltered pool (prefer ROJO + La Guaira)."""
    keys = [k for k in (row["nombre_norm"], row["nombre_archivo_norm"]) if k]
    tokens = {tok for k in keys for tok in k.split() if len(tok) >= 4}
    if not keys and not (row["lat_informe"] and row["lon_informe"]):
        return []

    # narrow by token presence when possible
    pool = hab_pool
    if tokens:
        mask = pd.Series(False, index=pool.index)
        for tok in tokens:
            mask = mask | pool["_norm_nombre"].str.contains(re.escape(tok), na=False)
            mask = mask | pool["_norm_dir"].str.contains(re.escape(tok), na=False)
        narrowed = pool[mask]
        if len(narrowed) >= 1:
            pool = narrowed
        elif len(pool) > 8000:
            if row["lat_informe"] and row["lon_informe"]:
                try:
                    la, lo = float(row["lat_informe"]), float(row["lon_informe"])
                    pool = pool[
                        (pool["lat"].sub(la).abs() < 0.02) & (pool["lng"].sub(lo).abs() < 0.02)
                    ]
                except Exception:
                    pool = pool.iloc[0:0]

    candidates = []
    for _, h in pool.iterrows():
        score = _score_row(
            row,
            h["_norm_nombre"],
            h["_norm_dir"],
            h.get("lat"),
            h.get("lng"),
            keys,
            tokens,
        )
        if score >= 0.35:
            candidates.append((score, h))

    candidates.sort(key=lambda x: -x[0])
    out = []
    for score, h in candidates[:top_n]:
        out.append(
            {
                "score": round(float(score), 3),
                "id": h.get("id"),
                "certificado": str(h.get("certificado", "")).replace('="', "").replace('"', ""),
                "etiqueta": h.get("etiqueta"),
                "nombre": h.get("nombre_edificacion"),
                "direccion": h.get("direccion"),
                "municipio": h.get("municipio"),
                "parroquia": h.get("parroquia"),
                "num_pisos": h.get("num_pisos"),
                "material": h.get("material"),
                "uso": h.get("uso"),
                "inspector": h.get("inspector_nombre"),
                "created_at": h.get("created_at"),
                "estatus_demolicion": h.get("estatus_demolicion"),
                "lat": h.get("lat"),
                "lng": h.get("lng"),
                "observaciones": (str(h.get("observaciones") or "")[:220]),
            }
        )
    return out


def set_run_font(run, size=11, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def _abc_points(val, pts_a=0, pts_b=0, pts_c=0) -> float:
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return 0.0
    s = str(val).strip().upper()
    if s in ("C", "3"):
        return float(pts_c)
    if s in ("B", "2"):
        return float(pts_b)
    if s in ("A", "1"):
        return float(pts_a)
    return 0.0


def score_gravedad_rojo(row: pd.Series) -> dict:
    """Score 0–100 de gravedad / probabilidad relativa de demolición (Fase 1 Habitable).

    Heurística de priorización operativa — no sustituye dictamen de ingeniería.
    """
    detail = []
    total = 0.0

    def add(label: str, pts: float):
        nonlocal total
        if pts:
            total += pts
            detail.append(f"{label}+{pts:g}")

    # Riesgos de etiqueta ANIH (máx. 40)
    p = _abc_points(row.get("riesgo_externo"), 0, 8, 20)
    add("riesgo_externo", p)
    p = _abc_points(row.get("riesgo_severo"), 0, 8, 20)
    add("riesgo_severo", p)

    # Exterior / estructura (máx. 43)
    p = _abc_points(row.get("ext_colapso_estructura"), 0, 6, 15)
    add("ext_colapso", p)
    p = _abc_points(row.get("ext_peligro_aledanos"), 0, 4, 10)
    add("peligro_aledanos", p)
    p = _abc_points(row.get("ext_inclinacion"), 0, 3, 8)
    add("inclinacion", p)
    p = _abc_points(row.get("ext_asentamiento"), 0, 2, 5)
    add("asentamiento", p)
    p = _abc_points(row.get("ext_peligro_geologico"), 0, 2, 5)
    add("peligro_geologico", p)

    # Acciones recomendadas en planilla (máx. 8)
    acc = str(row.get("acc_medidas") or "").upper()
    if "DEMOLER" in acc:
        add("acc_demoler", 8)
    elif "APUNTALAR" in acc and "ACORDONAR" in acc:
        add("acc_apunt_acord", 4)
    elif "APUNTALAR" in acc or "ACORDONAR" in acc:
        add("acc_mitigacion", 2)

    # Gas (máx. 3)
    eg = row.get("emergencia_gas")
    if eg is True or str(eg).strip().lower() in ("true", "1", "si", "sí"):
        add("emergencia_gas", 3)

    # Consecuencia por altura (máx. 6)
    try:
        npisos = float(row.get("num_pisos"))
    except (TypeError, ValueError):
        npisos = float("nan")
    if pd.notna(npisos):
        if npisos >= 10:
            add("altura_10+", 6)
        elif npisos >= 6:
            add("altura_6-9", 5)
        elif npisos >= 3:
            add("altura_3-5", 3)
        elif npisos >= 1:
            add("altura_1-2", 1)

    # Piso crítico declarado (máx. 3)
    pc = row.get("piso_critico")
    if pc is not None and str(pc).strip() and str(pc).strip().upper() not in ("NAN", "NONE", ""):
        add("piso_critico", 3)

    # Conteos de elementos severos (máx. 5)
    sev_sum = 0.0
    for c in ("sev_columna", "sev_viga", "sev_muro_concreto", "sev_muro_mamposteria"):
        try:
            v = float(row.get(c))
            if pd.notna(v) and v > 0:
                sev_sum += min(v, 10)
        except (TypeError, ValueError):
            pass
    if sev_sum >= 8:
        add("sev_elementos", 5)
    elif sev_sum >= 3:
        add("sev_elementos", 3)
    elif sev_sum > 0:
        add("sev_elementos", 1)

    # Componentes / moderado en C (máx. 4)
    p = _abc_points(row.get("riesgo_componentes"), 0, 1, 2)
    add("riesgo_componentes", p)
    p = _abc_points(row.get("riesgo_moderado"), 0, 1, 2)
    add("riesgo_moderado", p)

    total = min(100.0, round(total, 1))

    if total >= 80:
        banda = "Muy alta"
        prob = "Muy alta probabilidad relativa de recomendación de demolición"
    elif total >= 60:
        banda = "Alta"
        prob = "Alta probabilidad relativa — priorizar verificación detallada"
    elif total >= 40:
        banda = "Media"
        prob = "Probabilidad media — verificar; puede ser ROJO por riesgo localizado"
    elif total >= 20:
        banda = "Baja-media"
        prob = "Menor señal de demolición en Fase 1 — no descartar sin revisión"
    else:
        banda = "Baja / datos incompletos"
        prob = "Poca señal en campos de gravedad o campos vacíos — curar datos"

    return {
        "score_gravedad": total,
        "banda_prioridad": banda,
        "prob_relativa_demolicion": prob,
        "score_detalle": "; ".join(detail) if detail else "sin aportes",
    }


def build_ranking_rojo(hab_rojo: pd.DataFrame, ids_en_informes: set) -> pd.DataFrame:
    scores = hab_rojo.apply(score_gravedad_rojo, axis=1, result_type="expand")
    out = hab_rojo.copy()
    out = pd.concat([out.reset_index(drop=True), scores.reset_index(drop=True)], axis=1)
    out["en_lote_informes_pdf"] = out["id"].isin(ids_en_informes).map({True: "Sí", False: "No"})
    out = out.sort_values(
        ["score_gravedad", "num_pisos", "id"],
        ascending=[False, False, True],
    ).reset_index(drop=True)
    out.insert(0, "rank_gravedad", out.index + 1)

    keep = [
        "rank_gravedad",
        "score_gravedad",
        "banda_prioridad",
        "prob_relativa_demolicion",
        "score_detalle",
        "en_lote_informes_pdf",
        "id",
        "certificado",
        "etiqueta",
        "nombre_edificacion",
        "direccion",
        "estado",
        "municipio",
        "parroquia",
        "uso",
        "material",
        "num_pisos",
        "piso_critico",
        "riesgo_externo",
        "riesgo_severo",
        "riesgo_moderado",
        "riesgo_componentes",
        "ext_colapso_estructura",
        "ext_peligro_aledanos",
        "ext_peligro_geologico",
        "ext_asentamiento",
        "ext_inclinacion",
        "emergencia_gas",
        "acc_medidas",
        "acc_inspecciones",
        "sev_columna",
        "sev_viga",
        "sev_muro_concreto",
        "sev_muro_mamposteria",
        "inspector_nombre",
        "ente",
        "created_at",
        "lat",
        "lng",
        "observaciones",
        "estatus_demolicion",
    ]
    keep = [c for c in keep if c in out.columns]
    return out[keep]


def build_indice_pestanas() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "pestana": "Indice",
                "contenido": "Guía del libro operativo.",
                "como_usarla": "Empezar aquí.",
            },
            {
                "pestana": "Resumen",
                "contenido": "Indicadores del lote PDF y del universo ROJO.",
                "como_usarla": "Vista ejecutiva rápida.",
            },
            {
                "pestana": "Cruce_informes",
                "contenido": "Informes detallados del lote PDF cruzados con Habitable.",
                "como_usarla": "Seguimiento del lote ya visitado en detalle.",
            },
            {
                "pestana": "Ranking_ROJO",
                "contenido": "Todos los ROJO con score, banda y flags de filtro "
                "(Es_La_Guaira, Es_Top200, Prioridad_visita, En_lote_PDF).",
                "como_usarla": "Autofiltro: Prioridad_visita=Priorizar; Es_La_Guaira=Sí; Es_Top200=Sí.",
            },
            {
                "pestana": "Control_2da_ronda",
                "contenido": "Cola de trabajo 2.ª visita: validación, decisión D1–D5, magnitud M.",
                "como_usarla": "Actualizar estado y dictamen; listas desplegables.",
            },
            {
                "pestana": "Catalogo_campos",
                "contenido": "Campos del formato de inspección detallada (A0–F).",
                "como_usarla": "Referencia de diseño / sistema.",
            },
            {
                "pestana": "(archivo aparte) Ejemplo vaciado… Franco Mar.xlsx",
                "contenido": "Vaciado digital tipo informe; acompaña el Word de formato.",
                "como_usarla": "Capacitación / propuesta metodológica.",
            },
            {
                "pestana": "Listas_desplegables",
                "contenido": "Catálogo de valores para validaciones (hoja oculta).",
                "como_usarla": "No editar salvo para ampliar opciones.",
            },
        ]
    )


def build_criterios_score_docx(path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        t.add_run(
            "CRITERIOS DEL SCORE DE GRAVEDAD (ETIQUETAS ROJAS)\n"
            "Guía didáctica · Priorización para verificación detallada\n"
            "y probabilidad relativa de recomendación de demolición"
        ),
        13,
        True,
    )
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            f"Documento de trabajo · {date.today().strftime('%d/%m/%Y')} · "
            "Basado en campos de la planilla Habitable (Fase 1 / ANIH). "
            "Versión ampliada con ejemplos numéricos."
        ),
        9,
    )

    doc.add_heading("1. Idea en una frase", level=1)
    doc.add_paragraph(
        "Cada inspección ya marcada ROJO recibe un número de 0 a 100. "
        "Ese número no dice «demoler» o «no demoler»: dice «qué tan urgente "
        "parece, con lo que ya llenó el inspector en Fase 1, revisar este caso "
        "en una segunda ronda detallada»."
    )
    doc.add_paragraph(
        "Piense el score como una nota de prioridad en una cola de emergencia: "
        "100 = atender primero; 20 = sigue siendo ROJO, pero con menos señales "
        "de colapso/demolición en los campos capturados."
    )

    doc.add_heading("2. Qué NO es este score", level=1)
    add_bullets(
        doc,
        [
            "No es un dictamen de ingeniería estructural.",
            "No es una orden de demolición ni un permiso municipal.",
            "No sustituye la inspección detallada ni el proyecto ejecutivo.",
            "Sí es un ranking operativo revisable para adelantar trabajo "
            "(cuadrillas, mesas técnicas, listados cortos).",
        ],
    )

    doc.add_heading("3. Cómo se lee la planilla (A / B / C)", level=1)
    doc.add_paragraph(
        "En Habitable, varios bloques de la ANIH usan letras. En este score "
        "las interpretamos así (simplificado para priorizar):"
    )
    add_bullets(
        doc,
        [
            "A — condición favorable / sin el daño o riesgo de ese ítem (aporta 0 o muy poco).",
            "B — condición intermedia (aporta puntos medios).",
            "C — condición adversa / daño o riesgo alto en ese ítem (aporta el máximo del ítem).",
            "Vacío (sin dato) — no suma. Por eso un ROJO «poco llenado» puede "
            "quedar con score bajo aunque el edificio sea grave en la realidad.",
        ],
    )
    doc.add_paragraph(
        "Regla práctica: dos letras C en los riesgos principales "
        "(externo + severo) ya empujan el caso hacia la zona Alta / Muy alta."
    )

    doc.add_heading("4. Tabla de puntos (máximo 100)", level=1)
    doc.add_paragraph(
        "Los aportes se suman. Si la suma supera 100, se deja en 100. "
        "En el Excel, la columna score_detalle muestra exactamente qué sumó "
        "(ejemplo: riesgo_externo+20; ext_colapso+15; altura_10++6)."
    )
    add_bullets(
        doc,
        [
            "riesgo_externo: B = 8 · C = 20  (peligro hacia afuera / entorno).",
            "riesgo_severo: B = 8 · C = 20  (daño severo estructural en planilla).",
            "ext_colapso_estructura: B = 6 · C = 15.",
            "ext_peligro_aledanos: B = 4 · C = 10.",
            "ext_inclinacion: B = 3 · C = 8.",
            "ext_asentamiento: B = 2 · C = 5.",
            "ext_peligro_geologico: B = 2 · C = 5.",
            "acc_medidas: si aparece DEMOLER → +8; si apuntalar y acordonar → +4; "
            "si solo una de esas → +2.",
            "emergencia_gas = sí → +3.",
            "Altura (num_pisos): 1–2 → +1; 3–5 → +3; 6–9 → +5; 10 o más → +6.",
            "piso_critico informado (cualquier valor) → +3.",
            "Conteos sev_* (columnas/vigas/muros): hasta +5 según cuántos elementos "
            "severos se registraron.",
            "riesgo_componentes / riesgo_moderado en C → hasta +2 cada uno.",
        ],
    )

    doc.add_heading("5. Bandas: qué significa el número", level=1)
    add_bullets(
        doc,
        [
            "80–100 · Muy alta — varias señales graves a la vez; máxima prioridad "
            "de verificación detallada; fuerte probabilidad relativa de que el "
            "dictamen posterior oriente a demolición.",
            "60–79 · Alta — priorizar en la cola de segunda ronda.",
            "40–59 · Media — verificar; puede ser ROJO por riesgo importante pero "
            "no necesariamente demolición inmediata.",
            "20–39 · Baja-media — menor señal acumulada en Fase 1; no excluir "
            "sin revisión (a veces falta llenado).",
            "0–19 · Baja / datos incompletos — pocos campos de gravedad marcados "
            "en C, o edificio alto con poca señal estructural en planilla.",
        ],
    )

    doc.add_heading("6. Ejemplos didácticos (cómo se arma el valor)", level=1)
    doc.add_paragraph(
        "Los ejemplos siguientes son ilustrativos: muestran la lógica aritmética. "
        "Los nombres y puntajes se alinean con el tipo de casos que aparecen en "
        "Ranking_ROJO_nacional del Excel (corte Habitable 20/08/2026)."
    )

    doc.add_heading("Ejemplo A — Score ≈ 86 (banda Muy alta)", level=2)
    doc.add_paragraph(
        "Perfil: edificio de varios pisos, riesgos externos y severos en C, "
        "colapso estructural en C, acción con DEMOLER y algo de peligro a vecinos."
    )
    add_bullets(
        doc,
        [
            "riesgo_externo C → +20",
            "riesgo_severo C → +20",
            "ext_colapso_estructura C → +15",
            "ext_peligro_aledanos B → +4",
            "ext_inclinacion B → +3",
            "ext_peligro_geologico B → +2",
            "acc_medidas con DEMOLER → +8",
            "altura 6–9 pisos → +5",
            "piso_critico informado → +3",
            "Suma ≈ 20+20+15+4+3+2+8+5+3 = 80 (puede llegar a 86 si hay más "
            "ítems B/C o sev_*).",
        ],
    )
    doc.add_paragraph(
        "Lectura: casi todos los «interruptores» graves están encendidos. "
        "En la cola, este caso va primero. La segunda ronda debe confirmar "
        "si procede demolición controlada."
    )

    doc.add_heading("Ejemplo B — Score ≈ 81 (Muy alta) aunque sea 1 piso", level=2)
    doc.add_paragraph(
        "Perfil: vivienda o edificación baja, pero con colapso + riesgos C + "
        "peligro aledaños C."
    )
    add_bullets(
        doc,
        [
            "riesgo_externo C → +20",
            "riesgo_severo C → +20",
            "ext_colapso C → +15",
            "peligro_aledanos C → +10",
            "inclinación B → +3",
            "asentamiento B → +2",
            "altura 1–2 → +1",
            "Suma ≈ 71; con DEMOLER (+8) u otros C llega a ~80+.",
        ],
    )
    doc.add_paragraph(
        "Lectura didáctica: la altura suma poco, pero el daño/riesgo suma mucho. "
        "Un caso de 1 piso puede superar a un edificio alto «poco documentado»."
    )

    doc.add_heading("Ejemplo C — Score ≈ 77 (banda Alta)", level=2)
    doc.add_paragraph(
        "Perfil: torre o bloque alto (p. ej. 14 pisos), riesgo externo C, colapso C, "
        "aledaños C, inclinación C, acción DEMOLER — pero riesgo_severo aún en A "
        "(no aporta)."
    )
    add_bullets(
        doc,
        [
            "riesgo_externo C → +20",
            "riesgo_severo A → +0",
            "ext_colapso C → +15",
            "peligro_aledanos C → +10",
            "inclinación C → +8",
            "asentamiento B → +2",
            "peligro geológico C → +5",
            "DEMOLER → +8",
            "altura 10+ → +6",
            "Suma ≈ 74–77.",
        ],
    )
    doc.add_paragraph(
        "Lectura: sigue siendo prioridad alta. Falta el «doble C» de riesgos "
        "ANIH, por eso no cruza fácilmente a 80+, pero el paquete exterior + "
        "DEMOLER + altura lo mantiene arriba en la cola."
    )

    doc.add_heading("Ejemplo D — Score ≈ 59 (banda Media)", level=2)
    doc.add_paragraph(
        "Perfil: casa de 1–2 pisos con varios exteriores en C, pero sin el "
        "paquete completo (p. ej. sin DEMOLER o sin riesgo_severo C)."
    )
    add_bullets(
        doc,
        [
            "riesgo_externo C → +20",
            "ext_colapso C → +15",
            "peligro_aledanos C → +10",
            "inclinación / asentamiento / geológico (mezcla B/C) → +8 a +13",
            "altura 1–2 → +1",
            "Suma típica ≈ 54–59.",
        ],
    )
    doc.add_paragraph(
        "Lectura: hay daño serio registrado, pero el score no alcanza la "
        "franja Alta. Conviene verificación; la demolición no es la única "
        "hipótesis (puede ser estabilización, remoción parcial, etc.)."
    )

    doc.add_heading("Ejemplo E — Score ≈ 39 (banda Baja-media) en torre alta", level=2)
    doc.add_paragraph(
        "Perfil: edificio de 15–16 pisos, riesgo_severo C y quizá DEMOLER, "
        "pero riesgo_externo A/B y sin colapso exterior en C."
    )
    add_bullets(
        doc,
        [
            "riesgo_severo C → +20",
            "DEMOLER → +8",
            "altura 10+ → +6",
            "piso_critico → +3",
            "algo de componentes/moderado → +1 a +2",
            "Suma ≈ 38–39.",
        ],
    )
    doc.add_paragraph(
        "Lectura didáctica clave: ser alto no basta. Sin C en externo/colapso/"
        "aledaños, el score se queda en Baja-media. Puede ser un ROJO «real» "
        "con planilla incompleta: por eso no se descarta, se revisa el llenado "
        "antes de bajarlo de la cola."
    )

    doc.add_heading("Ejemplo F — Score ≈ 19 (Baja / datos incompletos)", level=2)
    doc.add_paragraph(
        "Perfil: torre alta (15–18 pisos) con pocos C; a veces solo DEMOLER "
        "o un riesgo B, más altura y piso crítico."
    )
    add_bullets(
        doc,
        [
            "DEMOLER → +8",
            "altura 10+ → +6",
            "piso_critico → +3",
            "riesgo_componentes C → +2",
            "Suma ≈ 19.",
        ],
    )
    doc.add_paragraph(
        "Lectura: el sistema «no ve» gravedad estructural en los campos A/B/C. "
        "Puede ser subregistro. Acción sugerida: no usar este score para "
        "cerrar el caso; pedir completar planilla o inspección detallada "
        "si hay alerta territorial."
    )

    doc.add_heading("7. Comparación rápida (para interiorizar la escala)", level=1)
    add_bullets(
        doc,
        [
            "Casa baja + colapso C + externo C + aledaños C → puede ser Muy alta (~80).",
            "Torre 14 pisos + externo C + colapso C + DEMOLER + inclinación C → Alta (~75).",
            "Casa con varios C exteriores pero sin DEMOLER/severo C → Media (~55–59).",
            "Torre alta + severo C + DEMOLER pero externo A y sin colapso C → Baja-media (~39).",
            "Torre alta casi sin C → Baja / incompleta (~15–20).",
        ],
    )

    doc.add_heading("8. Cómo usarlo mañana en el Excel", level=1)
    add_bullets(
        doc,
        [
            "Abrir Ranking_ROJO_LaGuaira (o nacional).",
            "Filtrar banda_prioridad = Muy alta o Alta (score ≥ 60) para la primera ola.",
            "Mirar score_detalle para explicar en mesa técnica por qué subió.",
            "Usar en_lote_informes_pdf = Sí para ver si ese id ya tiene informe del lote actual.",
            "Bajar a Media solo cuando la ola Alta esté cubierta o haya criterio territorial.",
            "Tratar Baja / incompleta como «curar dato», no como «edificio sano».",
        ],
    )

    doc.add_heading("9. Limitaciones (para no sobreinterpretar)", level=1)
    add_bullets(
        doc,
        [
            "Depende de la calidad del llenado en Fase 1.",
            "No analiza fotos ni el texto de observaciones de forma semántica.",
            "estatus_demolicion en el export puede venir vacío aunque el caso ya tenga informe PDF.",
            "Cuando existan dictámenes detallados sí/no demolición, el score debe recalibrarse "
            "(pesos A/B/C y umbrales de banda).",
        ],
    )

    doc.add_heading("10. Cierre", level=1)
    doc.add_paragraph(
        "El valor del score es una suma transparente de señales ya capturadas. "
        "Úselo para ordenar el trabajo; valide siempre con inspección detallada "
        "antes de cualquier decisión de demolición."
    )

    doc.save(str(path))


def _add_table(doc, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = "Calibri"
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.name = "Calibri"
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph("")
    return table


def build_formato_inspeccion_docx(path: Path):
    """Guía didáctica completa + precarga/ranking + anexo Franco Mar."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        t.add_run("COMISIÓN PRESIDENCIAL PARA LA EVALUACIÓN DE HABITABILIDAD"),
        12,
        True,
    )
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        t2.add_run(
            "GUÍA DEL FORMATO DE INSPECCIÓN DETALLADA\n"
            "(Segunda visita a edificios con etiqueta ROJA)\n"
            "Precarga de la inspección previa + ranking de prioridad + "
            "decisión demoler / reparar"
        ),
        13,
        True,
    )
    meta = doc.add_paragraph()
    set_run_font(
        meta.add_run(
            f"Versión didáctica ampliada · {date.today().strftime('%d/%m/%Y')} · "
            "Incluye explicaciones por sección y Anexo 1 (Franco Mar) con "
            "inspección previa Habitable, ranking y visita detallada."
        ),
        9,
    )

    # ------------------------------------------------------------------
    doc.add_heading("1. ¿Qué problema resuelve este formato?", level=1)
    doc.add_paragraph(
        "Hoy ya existen muchos informes buenos, pero cada uno se escribe distinto. "
        "Eso dificulta saber, en una sola lista: cuántos van a demolición, cuántos "
        "a reparación y qué tan grande es esa reparación."
    )
    doc.add_paragraph(
        "Además, la primera inspección (Habitable / planilla rápida) ya dejó datos "
        "y un ranking de gravedad para priorizar visitas. Este segundo informe "
        "no parte de cero: recibe esa información, la valida o corrige, profundiza "
        "el daño y cierra una decisión de control."
    )
    doc.add_paragraph(
        "En conjunto, el segundo informe debe poder responder con claridad:"
    )
    add_bullets(
        doc,
        [
            "1) ¿Confirmamos que sigue siendo ROJO / inhabitable?",
            "2) ¿La decisión es demoler, reparar/reconstruir, pedir más estudios, "
            "tratar escombros, o solo vigilar?",
            "3) Si se repara: ¿la intervención es pequeña, mediana, grande o casi total?",
            "4) ¿Qué traía la inspección previa y el ranking, y qué se corrigió en campo?",
            "5) ¿Qué medidas inmediatas hay que mantener (perímetro, desalojo, etc.)?",
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("2. Dos tipos de información en el formato", level=1)
    doc.add_paragraph(
        "El formato combina campos estructurados (comparables entre edificaciones y "
        "aptos para tableros o sistemas) con secciones de análisis libre, donde el "
        "ingeniero documenta su criterio técnico con libertad profesional."
    )
    _add_table(
        doc,
        ["Tipo", "Descripción", "Ejemplo"],
        [
            [
                "Casillas cerradas",
                "Respuestas cortas, comparables entre edificios",
                "14 pisos · Demoler · Prioridad inmediata",
            ],
            [
                "Semi-cerradas (escalas)",
                "A/B/C o grados de daño con definición escrita",
                "Riesgo fachada C · columnas grado IV",
            ],
            [
                "Referencia precargada",
                "Viene de Habitable / ranking; se valida",
                "ID 171818 · score 51 · banda Media",
            ],
            [
                "Texto libre del ingeniero",
                "Explicación técnica con libertad profesional",
                "«Fallaron las columnas del nivel 1…»",
            ],
            [
                "Fotos / croquis",
                "Prueba visual de lo dicho",
                "Foto del piso crítico + etiqueta",
            ],
        ],
        [3.5, 5.5, 6.5],
    )

    # ------------------------------------------------------------------
    doc.add_heading("3. Cómo se parece a los informes que ya tienen", level=1)
    doc.add_paragraph(
        "No se descarta lo ya elaborado. Se homologa al mismo molde de captura:"
    )
    _add_table(
        doc,
        ["Tipo de informe actual", "Qué aporta", "Dónde cae en este formato"],
        [
            [
                "Ficha PNUD / ATC-20 (p. ej. OPPPE30)",
                "GPS, geometría, % columnas, inclinación, riesgos A/B/C",
                "Precarga + Partes B y C",
            ],
            [
                "Informe técnico nivel 2 (p. ej. Capri / Franco Mar)",
                "Piso crítico, falla por compresión, inviabilidad de reparación, fotos",
                "Partes C y D + textos libres",
            ],
            [
                "Evaluación patológica (p. ej. Arenal)",
                "Daño en vigas/losas/cerramientos y dictamen",
                "Matriz C + decisión D",
            ],
            [
                "Justificación corta (p. ej. Samantha, Brisamar)",
                "Conclusión demoler + firmas",
                "Al menos A0/A + D + firmas; conviene completar C",
            ],
            [
                "Recolección de escombros (p. ej. Fioremar)",
                "Ya colapsó; volumen a retirar",
                "Decisión D4 + magnitud de remoción",
            ],
        ],
        [4.2, 5.5, 6.0],
    )

    # ------------------------------------------------------------------
    doc.add_heading("4. Mapa mental: las partes del segundo informe", level=1)
    add_bullets(
        doc,
        [
            "A0 — Precarga: inspección previa Habitable + ranking (validar / corregir).",
            "A — Identidad confirmada del edificio.",
            "B — Cómo es el edificio (pisos, uso, estructura, ocupación).",
            "C — Qué daño se vio en detalle + texto libre del ingeniero.",
            "D — Decisión demoler / reparar / … + tamaño si repara + justificación libre.",
            "E — Fotos y firmas.",
            "Resumen ejecutivo — párrafo corto para mesa de trabajo.",
        ],
    )
    doc.add_paragraph(
        "Abajo, cada parte se explica con: para qué sirve, qué preguntas contesta "
        "y qué datos conviene llenar."
    )

    # ----- A0 -----
    doc.add_heading("5. Parte A0 — Precarga de la inspección previa + ranking", level=1)
    doc.add_paragraph(
        "El inspector de la segunda visita no parte de una hoja en blanco. "
        "El sistema (o el Excel de trabajo) le entrega lo ya conocido de la "
        "inspección rápida y el ranking que ordena la cola de visitas."
    )
    doc.add_paragraph("Preguntas que contesta esta parte:")
    add_bullets(
        doc,
        [
            "¿Qué dijo Habitable de este edificio (etiqueta, riesgos, acciones)?",
            "¿Qué tan prioritario era según el score (0–100) y su banda?",
            "¿En qué puesto de la cola nacional / La Guaira estaba?",
            "¿El inspector confirma, corrige o rechaza esos datos al llegar al sitio?",
        ],
    )
    doc.add_paragraph("Qué suele venir ya cargado:")
    add_bullets(
        doc,
        [
            "ID / certificado Habitable, nombre, dirección, GPS, municipio/parroquia.",
            "Etiqueta ROJA, fecha e inspector de Fase 1, riesgos A/B/C, piso crítico, acciones.",
            "Ranking: score, banda (Muy alta / Alta / Media…), puesto, detalle de puntos, "
            "texto de probabilidad relativa.",
            "Observaciones cortas de la primera inspección (si existen).",
        ],
    )
    doc.add_paragraph(
        "Importante: el ranking prioriza; no es el dictamen final. Un caso en banda "
        "Media puede terminar en demolición tras la visita detallada (ver Anexo 1)."
    )
    _add_table(
        doc,
        ["Acción del inspector", "En la práctica"],
        [
            ["Validar", "Marca «Correcto» si nombre, pisos, dirección y riesgos coinciden"],
            ["Corregir", "Si había errores (pisos, GPS, sótanos…), los enmenda en la ficha"],
            ["Confirmar etiqueta", "¿Sigue ROJO / inhabitable? Sí / No / Aún no se puede decir"],
            ["Usar el ranking como guía", "Sirve de referencia de prioridad; el dictamen lo define la visita detallada"],
        ],
        [4.5, 11.5],
    )

    # ----- A -----
    doc.add_heading("6. Parte A — Identificación confirmada", level=1)
    doc.add_paragraph(
        "Después de validar la precarga, se deja fija la identidad del caso para "
        "trazabilidad (mismo certificado Habitable, mismo edificio)."
    )
    doc.add_paragraph("Preguntas que contesta:")
    add_bullets(
        doc,
        [
            "¿De qué edificio exacto estamos hablando (nombre oficial + alias)?",
            "¿Dónde está (dirección, GPS de control de esta visita)?",
            "¿Quién inspeccionó en la segunda visita y quién supervisa?",
            "¿Cuándo se hizo la visita detallada?",
        ],
    )
    doc.add_paragraph("Datos típicos: nombre, dirección, GPS de control, fecha/hora, "
                      "equipo evaluador (C.I., CIV, ente), supervisor, vínculo al certificado.")

    # ----- B -----
    doc.add_heading("7. Parte B — Cómo es el edificio", level=1)
    doc.add_paragraph(
        "Describe la «ficha técnica» del inmueble. Sirve para entender la "
        "consecuencia del daño (no es lo mismo 2 pisos que 12) y para comparar casos."
    )
    doc.add_paragraph("Preguntas que contesta:")
    add_bullets(
        doc,
        [
            "¿Para qué se usa (vivienda, hotel, comercio, en construcción…)?",
            "¿Cuántos pisos, semisótanos y sótanos tiene realmente?",
            "¿De qué está hecho el sistema que lo sostiene (pórticos, muros, mixto…)?",
            "¿Está desalojado, con gente, o ya en escombros?",
            "¿Hay vecinos o vías que se verían afectados si colapsa?",
        ],
    )

    # ----- C -----
    doc.add_heading("8. Parte C — Qué daño se vio (matriz + texto libre)", level=1)
    doc.add_paragraph(
        "Aquí se homologa lo que ya escriben los informes tipo Capri, Franco Mar, "
        "OPPPE o Arenal: piso crítico, porcentaje de columnas graves, inclinación, "
        "vigas, peligro a vecinos. Las casillas permiten comparar; el texto libre "
        "permite al ingeniero explicar el mecanismo del daño y lo que no pudo ver."
    )
    doc.add_paragraph("Preguntas que contesta:")
    add_bullets(
        doc,
        [
            "¿Cuál es el piso / entrepiso más crítico?",
            "¿Qué tan afectadas están las columnas (exteriores e interiores)?",
            "¿Hay daño relevante en vigas, losas, muros o núcleo de escaleras?",
            "¿El edificio perdió verticalidad / se inclinó? ¿Se midió?",
            "¿Los elementos no estructurales (fachada, tanques, gas) agregan riesgo?",
            "¿Hay peligro claro para aledaños o la vía pública?",
            "¿Qué interpreta el ingeniero (causa probable, límites de la inspección visual)?",
        ],
    )
    doc.add_paragraph("Casillas sugeridas (cerradas / semi):")
    add_bullets(
        doc,
        [
            "Piso crítico (nivel).",
            "% aproximado de columnas en daño grave (grados III / IV / V o equivalente).",
            "Inclinación: sí/no + Δ (cm por metro) si se midió.",
            "Riesgos no estructurales A/B/C (fachada, escaleras, ascensores, tanques, gas).",
            "Peligro aledaños / vía: sí/no (+ distancia aproximada).",
        ],
    )
    doc.add_paragraph(
        "Caja libre (obligatoria, al menos un párrafo): análisis del evaluador — "
        "mecanismo probable, incertidumbre, disenso entre ingenieros si lo hay."
    )

    # ----- D -----
    doc.add_heading("9. Parte D — Decisión de control (el corazón del informe)", level=1)
    doc.add_paragraph(
        "Esta es la salida que permite gobernar el proceso: no basta con narrar el daño; "
        "hay que elegir un código de decisión comparable entre todos los casos ROJOS."
    )
    doc.add_paragraph("Preguntas que contesta:")
    add_bullets(
        doc,
        [
            "¿Demolemos, reparamos, pedimos más estudios, son escombros, o solo vigilamos?",
            "Si reparamos: ¿qué tan grande es la intervención (M1–M4)?",
            "¿Es viable reparar en sitio con seguridad, o no?",
            "¿Qué prioridad operativa tiene (inmediata / alta / programable)?",
            "¿Qué medidas inmediatas hay que mantener?",
            "¿Por qué el ingeniero eligió esa decisión (texto libre)?",
        ],
    )
    _add_table(
        doc,
        ["Código", "Decisión (lenguaje simple)", "¿Qué sigue?"],
        [
            ["D1", "Se recomienda DEMOLER", "Proyecto de demolición controlada + perímetro"],
            ["D2", "Se recomienda REPARAR / RECONSTRUIR", "Obligatorio marcar tamaño M1–M4"],
            ["D3", "Hacen falta MÁS ESTUDIOS", "No decidir demoler ni reparar todavía"],
            ["D4", "Ya está en ESCOMBROS / colapsó", "Control de remoción (no es demoler torre en pie)"],
            ["D5", "Seguir inhabitado y VIGILAR", "Perímetro + monitoreo; sin demolición aún"],
        ],
        [1.8, 5.5, 8.0],
    )

    doc.add_heading("9.1 Si es D2: tamaño de la reparación (M1–M4)", level=2)
    doc.add_paragraph(
        "Cuando no se demuele, el control necesita saber si hablamos de un arreglo "
        "local o de una obra enorme. Escala cerrada propuesta:"
    )
    _add_table(
        doc,
        ["Código", "Tamaño", "Definición operativa"],
        [
            ["M1", "Reparación pequeña / local", "Intervención en elementos puntuales; obra acotada"],
            ["M2", "Reparación importante", "Varios elementos o un nivel; requiere apuntalamiento"],
            ["M3", "Reconstrucción parcial", "Rehacer zonas amplias (p. ej. un piso crítico)"],
            ["M4", "Reconstrucción / refuerzo mayor", "Intervención sobre gran parte del sistema resistente"],
        ],
        [1.8, 4.2, 9.5],
    )
    doc.add_paragraph(
        "Campos adicionales útiles si D2: % aproximado de planta afectada; niveles a "
        "intervenir; ¿requiere refuerzo sísmico global?; plazo orientativo. "
        "Y siempre una caja libre: «por qué no demolición» / alcance de la reparación."
    )

    doc.add_heading("9.2 Si es D1: demolición", level=2)
    add_bullets(
        doc,
        [
            "Marcar inviabilidad de reparación en sitio (sí).",
            "Motivos tipificados (multi-selección): piso crítico masivo; pérdida de verticalidad; "
            "peligro aledaños; alto % de columnas graves; otro.",
            "Prioridad y medidas inmediatas (acordonar, excluir, apuntalar temporal, monitorear vecinos).",
            "Texto libre de justificación (como en Capri / Franco Mar / Samantha).",
        ],
    )

    doc.add_heading("9.3 Contraste rápido: ¿más demoler o más reparar?", level=2)
    _add_table(
        doc,
        ["Señal en campo", "Más compatible con D1 (demoler)", "Más compatible con D2 (reparar)"],
        [
            [
                "Piso crítico",
                "Aplastamiento / gran % de columnas graves en un nivel",
                "Daño concentrado reparable; bases inspeccionables",
            ],
            [
                "Verticalidad",
                "Pérdida de plomo / torre desalineada",
                "Sin desplome global relevante",
            ],
            [
                "Reparar en sitio",
                "Inviable o riesgo inaceptable para operarios",
                "Viabilidad técnica argumentada",
            ],
            [
                "Aledaños",
                "Alto riesgo si cae de forma descontrolada",
                "Riesgo controlable con perímetro / apuntalamiento",
            ],
            [
                "Magnitud M",
                "No aplica",
                "Obligatorio M1–M4 + texto de alcance",
            ],
        ],
        [3.5, 6.0, 6.0],
    )
    doc.add_paragraph(
        "Son orientaciones. El ingeniero puede discrepar: por eso existe el texto libre. "
        "Lo que no puede faltar es el código D1–D5 (y M1–M4 si D2)."
    )

    # ----- E -----
    doc.add_heading("10. Parte E — Evidencia y firmas", level=1)
    doc.add_paragraph("Preguntas que contesta:")
    add_bullets(
        doc,
        [
            "¿Hay prueba fotográfica de lo afirmado?",
            "¿Quién elaboró, revisó y aprobó (CIV)?",
        ],
    )
    add_bullets(
        doc,
        [
            "Mínimo sugerido: 6 fotos con descripción (vista general, piso crítico, columnas, "
            "vigas/losas, aledaños, etiqueta).",
            "Croquis de planta con marca del piso crítico (recomendado).",
            "Firmas: elaboró / revisó / aprobó.",
        ],
    )

    # ----- resumen + libertad -----
    doc.add_heading("11. Resumen ejecutivo (salida corta)", level=1)
    doc.add_paragraph(
        "Pregunta que contesta: ¿qué le digo a la mesa técnica en un párrafo? "
        "Al cierre, 8–12 líneas: edificio, qué traía Fase 1/ranking, qué se validó, "
        "qué se vio, decisión D (y M si aplica), prioridad y próxima acción."
    )

    doc.add_heading("12. Libertad del ingeniero (dónde escribe con libertad)", level=1)
    doc.add_paragraph(
        "Los códigos estructurados organizan el control y la comparabilidad entre casos; "
        "el texto libre documenta el análisis de ingeniería."
    )
    add_bullets(
        doc,
        [
            "En C: mecanismo del daño, límites de lo observado, dudas.",
            "En D: por qué demoler o reparar; por qué el ranking no basta solo.",
            "Observaciones abiertas: acceso impedido, seguridad del equipo, hallazgos no tipificados.",
            "Disenso: si hay discrepancia entre evaluadores, se registra y el supervisor fija el código final.",
        ],
    )

    doc.add_heading("13. Tabla ejemplo — datos que se pueden levantar", level=1)
    doc.add_paragraph(
        "Vista rápida de cómo se ve una ficha (mezcla de precarga + visita). "
        "El Anexo 1 desarrolla un caso real completo."
    )
    _add_table(
        doc,
        ["Campo", "Ejemplo", "Para qué sirve"],
        [
            ["ID / certificado Habitable", "171818 / 73420…", "Trazabilidad Fase 1"],
            ["Score / banda / puesto ranking", "51 / Media / ~115 La Guaira", "Prioridad de visita"],
            ["Validación precarga", "ROJO confirmado; corrige sótanos", "Calidad del dato"],
            ["Pisos confirmados", "12 + 2 sótanos", "Consecuencia"],
            ["Piso crítico", "Nivel 1", "Foco de daño"],
            ["% columnas graves", ">50% en piso crítico", "Severidad"],
            ["Decisión", "D1 demoler", "Control maestro"],
            ["Magnitud M", "N/A (si fuera D2: M2/M3…)", "Alcance si se repara"],
            ["Texto libre", "Párrafo del ingeniero", "Juicio profesional"],
            ["Fotos", "6+", "Evidencia"],
        ],
        [5.0, 5.5, 5.5],
    )

    doc.add_heading("14. Niveles para el sistema digital (A–F)", level=1)
    add_bullets(
        doc,
        [
            "Nivel A — Identidad Habitable + precarga ranking (bloquea inicio sin id/certificado).",
            "Nivel B — Geometría y configuración.",
            "Nivel C — Matriz de daño + texto libre.",
            "Nivel D — Decisión D1–D5 (+ M1–M4 si D2) + medidas + justificación.",
            "Nivel E — Evidencia y firmas.",
            "Nivel F — Estado del flujo (borrador / revisado / aprobado / publicado).",
        ],
    )

    doc.add_heading("15. Nota final", level=1)
    doc.add_paragraph(
        "Este formato no reemplaza el proyecto ejecutivo de demolición ni el diseño "
        "de refuerzo. Estandariza la segunda ronda para que demolición y reconstrucción "
        "queden como estados controlables, usando lo ya capturado en Fase 1 y el ranking "
        "como punto de partida. El Anexo 1 muestra el flujo completo con un caso real."
    )

    # ==================================================================
    # ANEXO 1 — Franco Mar (mantener / enriquecer un poco el puente)
    # ==================================================================
    doc.add_page_break()
    an = doc.add_paragraph()
    an.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        an.add_run(
            "ANEXO 1 — EJEMPLO DE INFORME YA LLENADO\n"
            "Edificio Franco Mar (La Guaira)\n"
            "Con inspección previa Habitable + ranking de prioridad + visita detallada"
        ),
        13,
        True,
    )
    nota = doc.add_paragraph()
    set_run_font(
        nota.add_run(
            "Este anexo usa datos reales del export Habitable (corte 20/08/2026), "
            "del ranking de gravedad del Excel de trabajo y del informe técnico "
            "detallado del 11/08/2026. Es un modelo de cómo se verá la ficha "
            "cuando el sistema precargue la Fase 1."
        ),
        9,
    )

    doc.add_heading("A0 — Lo que el sistema ya le muestra al inspector (precarga)", level=1)
    doc.add_paragraph(
        "Bloque de solo lectura al abrir la ficha. El inspector no lo inventa: lo recibe. "
        "Contesta: ¿qué sabemos antes de llegar? ¿por qué este caso está en la cola?"
    )

    doc.add_heading("A0.1 Inspección previa (Habitable / Fase 1)", level=2)
    _add_table(
        doc,
        ["Dato precargado", "Valor que trae Habitable"],
        [
            ["ID Habitable", "171818"],
            ["Certificado", "73420200720261338"],
            ["Nombre en sistema", "Edif. Francomar"],
            ["Etiqueta Fase 1", "ROJO"],
            ["Fecha de la inspección previa", "20/07/2026 (registro 13:44)"],
            ["Inspector Fase 1", "Maria Mercedes Nieves Duarte"],
            ["Ente", "Comisión Nacional de Ingenieros — CPEHI"],
            ["Dirección en Habitable", "Calle Principal Boulevard Tanaguarenas. Caraballeda. La Guaira."],
            ["Estado / Municipio / Parroquia", "La Guaira / Vargas / Caraballeda"],
            ["Uso / material", "Vivienda Edificio / Concreto"],
            ["Pisos / sótanos (Fase 1)", "12 / 1"],
            ["Año construcción (Fase 1)", "1980"],
            ["GPS Habitable", "10.6123543 , -66.8318071"],
            ["Riesgo externo / severo", "C / A"],
            ["Colapso estructura (exterior)", "C"],
            ["Piso crítico declarado en Fase 1", "Planta baja, piso 1, 2, 3, 4"],
            ["Acciones sugeridas Fase 1", "Acordonar"],
            ["Observación Fase 1", "«Edificio a punto de colapso, se recomienda Desalojar»"],
        ],
        [5.5, 10.5],
    )

    doc.add_heading("A0.2 Ranking de prioridad (ejemplo real)", level=2)
    doc.add_paragraph(
        "Contesta: ¿qué tan urgente parecía este ROJO solo con la planilla Fase 1?"
    )
    _add_table(
        doc,
        ["Indicador de ranking", "Valor en este ejemplo"],
        [
            ["Score de gravedad (0–100)", "51"],
            ["Banda de prioridad", "Media"],
            ["Puesto nacional (entre todos los ROJO)", "264 de ~9.270"],
            ["Puesto en La Guaira / Vargas", "115 de ~4.247"],
            [
                "Probabilidad relativa (texto del ranking)",
                "Media — verificar; puede ser ROJO por riesgo localizado",
            ],
            [
                "Por qué sumó puntos (score_detalle)",
                "riesgo_externo+20; ext_colapso+15; acordonar+2; altura_10++6; "
                "piso_critico+3; sev_elementos+5",
            ],
            [
                "Lectura para el inspector",
                "No es de los peores del país en la planilla Fase 1, pero sí prioritario "
                "en el litoral. La visita detallada debe confirmar si demoler o reparar.",
            ],
        ],
        [5.5, 10.5],
    )

    doc.add_heading("A0.3 Validación del inspector (casillas)", level=2)
    doc.add_paragraph(
        "Contesta: ¿qué de la precarga se sostiene y qué hay que corregir?"
    )
    _add_table(
        doc,
        ["Pregunta de validación", "Respuesta en este ejemplo (tras la visita)"],
        [
            ["¿El edificio precargado es el correcto?", "Sí — es Franco Mar / Francomar"],
            ["¿La etiqueta ROJA se mantiene?", "Sí — confirmada"],
            [
                "¿Los pisos de Fase 1 son correctos?",
                "Parcial — se confirman 12 pisos; en detalle se observan 2 sótanos "
                "(Fase 1 decía 1)",
            ],
            [
                "¿El GPS / dirección son útiles?",
                "Sí, con matiz — la visita usó también Av. José María España / Tanaguarena; "
                "GPS de control de campo distinto al de Habitable",
            ],
            [
                "¿El ranking Media refleja lo que vio?",
                "El ranking ayudó a priorizar; el daño detallado es más grave de lo que "
                "sugiere solo el score 51",
            ],
            [
                "¿Corrige algún dato de Fase 1?",
                "Sí — sótanos 2; precisar dirección/GPS de control; reforzar acciones "
                "(no solo acordonar)",
            ],
        ],
        [6.0, 10.0],
    )

    doc.add_heading("A/B — Identidad y edificio confirmados en la visita detallada", level=1)
    doc.add_paragraph(
        "Contesta: ¿quién es el edificio «en limpio» después de validar?"
    )
    _add_table(
        doc,
        ["Dato", "Valor tras validar / completar"],
        [
            ["Nombre confirmado", "Edificio Franco Mar"],
            ["Fecha visita detallada", "11 de agosto de 2026"],
            ["Código informe detallado", "DGPS_INF-F2-Tanaguarena /11/08/2026-001"],
            ["GPS de control (visita)", "Lat 10.489497 · Lon -66.899001"],
            ["Pisos / sótanos confirmados", "12 pisos + 2 sótanos (estacionamiento)"],
            ["Área planta aprox.", "780 m²"],
            [
                "Sistema",
                "Pórticos de concreto armado (~22 columnas perimetrales + 18 interiores)",
            ],
            ["Evaluadores visita 2", "Ing. Luis Burgos; Ing. José García"],
            ["Supervisor", "Ing. Aura Quintero"],
        ],
        [5.5, 10.5],
    )

    doc.add_heading("C — Daño en la visita detallada (casillas)", level=1)
    doc.add_paragraph(
        "Contesta: ¿qué tan grave está, con el detalle que Fase 1 no alcanza a mostrar?"
    )
    _add_table(
        doc,
        ["Dato", "Valor"],
        [
            ["Piso crítico", "Nivel 1 / planta baja de transición"],
            ["Columnas en piso crítico", "Más del 50% con falla grave / aplastamiento"],
            ["Verticalidad", "Pérdida visible de plomo / inclinación"],
            ["Vigas", "Grietas fuertes por corte en hasta 5 niveles"],
            ["Sótanos", "Daño en columnas y conexiones de escalera"],
            ["Peligro a vecinos / vía", "Sí — alto"],
            ["Riesgo declarado en visita 2", "Colapso inminente / inhabitable"],
        ],
        [5.0, 11.0],
    )

    doc.add_heading("C — Texto libre del ingeniero (ejemplo)", level=1)
    doc.add_paragraph(
        "«Se confirma el carácter crítico del Nivel 1: gran parte de las columnas "
        "perimetrales e interiores fallaron por compresión, bajó el entrepiso y la "
        "torre perdió verticalidad. Hay daño en vigas en varios niveles y afectación "
        "en sótanos. Respecto a la precarga Habitable: la etiqueta ROJA y el aviso "
        "de desalojo eran correctos; el ranking Media (51) no alcanza a reflejar "
        "la gravedad vista en sitio. No se considera viable una reparación segura "
        "bajo 12 pisos con los apoyos del piso crítico en ese estado.»"
    )

    doc.add_heading("D — Decisión de control (cierre)", level=1)
    doc.add_paragraph(
        "Contesta: ¿demoler o reparar? ¿qué prioridad? ¿qué medidas?"
    )
    _add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Decisión", "D1 — Se recomienda DEMOLER"],
            ["Tamaño de reparación M1–M4", "No aplica (no es D2)"],
            [
                "¿La precarga/ranking cambian la decisión?",
                "No — la visita detalla y confirma demolición",
            ],
            ["Prioridad operativa", "Inmediata"],
            [
                "Medidas",
                "Exclusión total, acordonar, no ingreso; preparación de demolición "
                "controlada; monitoreo de vecinos",
            ],
        ],
        [5.5, 10.5],
    )

    doc.add_heading("D — Justificación libre (ejemplo)", level=1)
    doc.add_paragraph(
        "«Aunque el ranking de Fase 1 ubicaba el caso en banda Media, la inspección "
        "detallada muestra piso crítico masivo, pérdida de verticalidad y peligro "
        "a colindantes. Se recomienda demolición controlada. El valor del ranking "
        "fue traer el caso a la cola de verificación; el valor de esta ficha es "
        "cerrar el control D1 con evidencia.»"
    )

    doc.add_heading("E — Evidencia y firmas", level=1)
    add_bullets(
        doc,
        [
            "Fotos del informe detallado: piso crítico, columnas, verticalidad, entorno.",
            "Elaboraron: equipo de la visita 11/08/2026.",
            "Supervisión: Ing. Aura Quintero.",
            "Queda vinculado el certificado Habitable 73420200720261338 / id 171818.",
        ],
    )

    doc.add_heading("Resumen ejecutivo (para mesa de trabajo)", level=1)
    doc.add_paragraph(
        "Franco Mar (Caraballeda, La Guaira). Fase 1: ROJO (20/07/2026), score 51 "
        "(banda Media; puesto ~115 en La Guaira). Visita detallada 11/08/2026: "
        "12 pisos + 2 sótanos, piso crítico en Nivel 1 con >50% columnas graves, "
        "pérdida de verticalidad, riesgo a vecinos. Validación: se confirma ROJO; "
        "se corrigen sótanos y se precisa GPS/dirección de control. "
        "Decisión: D1 demoler, prioridad inmediata. "
        "Aprendizaje: el ranking prioriza; la segunda visita decide."
    )

    doc.add_heading("Cómo se usaría esto el día de campo (paso a paso)", level=1)
    add_bullets(
        doc,
        [
            "1) Abrir ficha → ver A0 (Habitable + ranking).",
            "2) Llegar al sitio → validar o corregir casillas de A0.3.",
            "3) Completar daño (C) y fotos (E).",
            "4) Elegir D1–D5 (y M1–M4 si repara).",
            "5) Escribir el párrafo ejecutivo.",
            "6) Firmar → el caso sale de la cola de «pendiente de verificación».",
        ],
    )

    doc.add_heading("Mini-contraste: ¿y si fuera reparación?", level=1)
    doc.add_paragraph(
        "Para enseñar D2/M (no es el caso Franco Mar): si el daño fuera grave pero "
        "localizado, sin pérdida de plomo de toda la torre y con reparación viable, "
        "la ficha cerraría D2 + M2 o M3, con texto libre explicando por qué no demoler. "
        "Franco Mar, con la evidencia real, cierra en D1."
    )

    doc.save(str(path))


def build_esquema_sistema_docx(path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        t.add_run(
            "ESQUEMA GENERAL — SISTEMA DE VACÍADO / COLABORACIÓN\n"
            "Inspecciones detalladas de verificación (ROJO → demolición)"
        ),
        14,
        True,
    )

    p = doc.add_paragraph()
    set_run_font(
        p.add_run(f"Propuesta de arquitectura funcional · {date.today().strftime('%d/%m/%Y')}"),
        9,
    )

    doc.add_heading("1. Objetivo del sistema", level=1)
    doc.add_paragraph(
        "Permitir que equipos de campo e ingenieros supervisores capturen, validen y "
        "consulten la inspección detallada de segunda ronda, enlazada al expediente Habitable "
        "(certificado / etiqueta ROJA), reduciendo el vaciado manual desde PDF heterogéneos "
        "y acelerando el seguimiento de demolición."
    )

    doc.add_heading("2. Actores", level=1)
    add_bullets(
        doc,
        [
            "Inspector / evaluador de campo (captura Nivel A–E).",
            "Supervisor institucional (revisión y firma).",
            "Mesa técnica / DGPS (aprobación de dictamen demolición).",
            "Coordinación de datos / BI (calidad, tablero, cruce con mart Habitable).",
            "Autoridad local / Protección Civil (consumo de estados y perímetros).",
        ],
    )

    doc.add_heading("3. Flujo de alto nivel", level=1)
    add_bullets(
        doc,
        [
            "1) Selección del universo: etiquetas ROJAS desde Habitable + prioridad territorial.",
            "2) Asignación de cuadrilla y ficha digital prellenada (identidad + GPS + etiqueta).",
            "3) Captura en campo (app o formulario web con modo offline) según niveles A–E.",
            "4) Carga de fotos y geo; validaciones (no cerrar sin dictamen ni evidencia mínima).",
            "5) Revisión supervisor → aprobación mesa técnica.",
            "6) Publicación de estado: demolición recomendada / en estudio / no demolición / escombros.",
            "7) Sincronización a tablero BI y a estatus de demolición en Habitable (cuando exista canal).",
            "8) Opcional: asistente de extracción desde PDF legado (OCR + curación humana).",
        ],
    )

    doc.add_heading("4. Componentes sugeridos", level=1)
    add_bullets(
        doc,
        [
            "Módulo de catálogo / cola de trabajo (casos ROJO pendientes de verificación).",
            "Formulario digital del formato de inspección (secciones = niveles A–F).",
            "Almacén de evidencias (objeto + metadatos).",
            "Motor de reglas mínimas de consistencia del dictamen.",
            "Panel de seguimiento ejecutivo (conteos, mapa, avance por municipio/parroquia).",
            "Conector Habitable (lectura de inspecciones; escritura de estatus cuando se autorice).",
            "Export Excel/PDF del informe estandarizado.",
            "Módulo de backfill: ingesta de PDF existentes → borrador de ficha para curación.",
        ],
    )

    doc.add_heading("5. Fases de implementación", level=1)
    add_bullets(
        doc,
        [
            "Fase 0 (inmediata): Excel maestro + Word de formato + cruce con export Habitable.",
            "Fase 1: Formulario web/Streamlit/Django con campos Nivel A–D; fotos controladas; export Excel.",
            "Fase 2: Integración lectura Habitable + mapa + workflow de firmas; PDF de salida homologado.",
            "Fase 3: App móvil offline, geocercas, OCR de informes legacy, API de estatus demolición.",
        ],
    )

    doc.add_heading("6. Principios de diseño", level=1)
    add_bullets(
        doc,
        [
            "Una ficha = un certificado Habitable (o un bloque/torre explícitamente separado).",
            "El dictamen de demolición es un campo controlado, no solo un párrafo libre.",
            "Toda cifra crítica admite «no observable» con motivo.",
            "Auditoría: quién cambió el dictamen y cuándo.",
            "Lenguaje ejecutivo en resumen; detalle técnico en secciones de daño.",
        ],
    )

    doc.add_heading("7. Relación con el ecosistema actual", level=1)
    doc.add_paragraph(
        "Habitable concentra la inspección rápida y la etiqueta. El BI Habitable / PDNA "
        "concentra el análisis agregado. Este sistema de segunda ronda se inserta entre ambos: "
        "produce el dato de decisión demolición/no demolición con evidencia, alimenta el estatus "
        "operativo y permite informes ejecutivos homogéneos a partir de lo que hoy llega como "
        "PDF diversos."
    )

    doc.save(str(path))


def main():
    print("Leyendo Habitable…")
    hab = pd.read_csv(CSV_HAB, low_memory=False)
    if "certificado" in hab.columns:
        hab["certificado"] = (
            hab["certificado"].astype(str).str.replace(r'^="?|"?$', "", regex=True)
        )
    hab["_norm_nombre"] = hab["nombre_edificacion"].map(norm_name)
    hab["_norm_dir"] = hab["direccion"].map(norm_name)
    hab_rojo = hab[hab["etiqueta"].astype(str).str.upper().str.contains("ROJO", na=False)].copy()
    hab_lg = hab[
        hab["estado"].astype(str).str.contains("Guaira|Vargas", case=False, na=False)
        | hab["municipio"].astype(str).str.contains("Vargas", case=False, na=False)
    ]
    hab_rojo_lg = hab_rojo[
        hab_rojo["estado"].astype(str).str.contains("Guaira|Vargas", case=False, na=False)
        | hab_rojo["municipio"].astype(str).str.contains("Vargas", case=False, na=False)
    ]
    # pool de matching: priorizar ROJO La Guaira; ampliar a todo ROJO y luego La Guaira
    print(
        f"Habitable filas={len(hab)} ROJO={len(hab_rojo)} "
        f"LaGuaira={len(hab_lg)} ROJO_LaGuaira={len(hab_rojo_lg)}"
    )

    rows = []
    pdfs = sorted(SRC_PDF.glob("*.pdf"))
    print(f"PDFs={len(pdfs)}")
    for f in pdfs:
        try:
            text, pages = extract_pdf_text(f)
            safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in f.stem)[:90]
            (EXTRACT / f"{safe}.txt").write_text(text, encoding="utf-8", errors="replace")
            fields = extract_fields(f.name, text, pages)
            # override manual por patrón de nombre de archivo
            manual_id = None
            stem_n = norm_name(f.stem)
            for pat, hid in MANUAL_HAB_ID.items():
                if norm_name(pat) in stem_n or norm_name(pat) in fields["nombre_norm"]:
                    manual_id = hid
                    break
            matches = best_hab_matches(fields, hab_rojo_lg)
            if not matches or (matches[0]["score"] < 0.55):
                matches2 = best_hab_matches(fields, hab_rojo)
                if matches2 and (not matches or matches2[0]["score"] > matches[0]["score"]):
                    matches = matches2
            if not matches or (matches[0]["score"] < 0.55):
                matches3 = best_hab_matches(fields, hab_lg)
                if matches3 and (not matches or matches3[0]["score"] > matches[0]["score"]):
                    matches = matches3
            if manual_id is not None:
                hit = hab[hab["id"] == manual_id]
                if len(hit):
                    h = hit.iloc[0]
                    matches = [
                        {
                            "score": 1.0,
                            "id": h.get("id"),
                            "certificado": str(h.get("certificado", "")).replace('="', "").replace('"', ""),
                            "etiqueta": h.get("etiqueta"),
                            "nombre": h.get("nombre_edificacion"),
                            "direccion": h.get("direccion"),
                            "municipio": h.get("municipio"),
                            "parroquia": h.get("parroquia"),
                            "num_pisos": h.get("num_pisos"),
                            "material": h.get("material"),
                            "uso": h.get("uso"),
                            "inspector": h.get("inspector_nombre"),
                            "created_at": h.get("created_at"),
                            "estatus_demolicion": h.get("estatus_demolicion"),
                            "lat": h.get("lat"),
                            "lng": h.get("lng"),
                            "observaciones": (str(h.get("observaciones") or "")[:220]),
                        }
                    ] + [m for m in matches if m.get("id") != manual_id][:2]
            rojo_m = [m for m in matches if str(m.get("etiqueta", "")).upper() == "ROJO"]
            best = rojo_m or matches
            m0 = best[0] if best else {}
            score = m0.get("score") or 0
            if manual_id is not None and m0.get("id") == manual_id:
                calidad = "Alto (alias manual)"
            elif score >= 0.75:
                calidad = "Alto"
            elif score >= 0.45:
                calidad = "Medio — revisar"
            elif m0:
                calidad = "Bajo — revisar"
            else:
                calidad = "Sin match"
            fields.update(
                {
                    "match_calidad": calidad,
                    "match_score": m0.get("score", ""),
                    "hab_id": m0.get("id", ""),
                    "hab_certificado": m0.get("certificado", ""),
                    "hab_etiqueta": m0.get("etiqueta", ""),
                    "hab_nombre": m0.get("nombre", ""),
                    "hab_direccion": m0.get("direccion", ""),
                    "hab_municipio": m0.get("municipio", ""),
                    "hab_parroquia": m0.get("parroquia", ""),
                    "hab_num_pisos": m0.get("num_pisos", ""),
                    "hab_material": m0.get("material", ""),
                    "hab_uso": m0.get("uso", ""),
                    "hab_inspector": m0.get("inspector", ""),
                    "hab_created_at": m0.get("created_at", ""),
                    "hab_estatus_demolicion": m0.get("estatus_demolicion", ""),
                    "hab_lat": m0.get("lat", ""),
                    "hab_lng": m0.get("lng", ""),
                    "hab_obs": m0.get("observaciones", ""),
                    "match_alternativos": " | ".join(
                        f"{m.get('nombre')} ({m.get('etiqueta')}, score={m.get('score')}, id={m.get('id')})"
                        for m in best[1:3]
                    ),
                }
            )
            rows.append(fields)
            print(f"OK {f.name[:48]} | {calidad} | {fields.get('hab_nombre', '')}")
        except Exception as e:
            print("ERR", f.name, e)
            rows.append({"archivo_pdf": f.name, "error": str(e)})

    df = pd.DataFrame(rows)
    ids_en_informes = set()
    if "hab_id" in df.columns:
        for v in df["hab_id"]:
            try:
                if pd.notna(v) and str(v).strip() not in ("", "nan"):
                    ids_en_informes.add(int(float(v)))
            except (TypeError, ValueError):
                pass

    preferred = [
        "archivo_pdf",
        "paginas_pdf",
        "tipo_informe",
        "codigo_documento",
        "nombre_edificacion_informe",
        "ubicacion_informe",
        "lat_informe",
        "lon_informe",
        "fecha_inspeccion_informe",
        "evaluadores",
        "supervisor",
        "dictamen_etiqueta",
        "recomienda_demolicion",
        "inviabilidad_reparacion",
        "uso",
        "anio_construccion_informe",
        "num_pisos_informe",
        "num_sotanos_informe",
        "area_planta_m2",
        "sistema_estructural",
        "piso_critico_mencionado",
        "hallazgos_flags",
        "n_evidencias_fotograficas_citadas",
        "extracto_conclusiones",
        "texto_disponible",
        "chars_texto_extraido",
        "match_calidad",
        "match_score",
        "hab_id",
        "hab_certificado",
        "hab_etiqueta",
        "hab_nombre",
        "hab_direccion",
        "hab_municipio",
        "hab_parroquia",
        "hab_num_pisos",
        "hab_material",
        "hab_uso",
        "hab_inspector",
        "hab_created_at",
        "hab_estatus_demolicion",
        "hab_lat",
        "hab_lng",
        "hab_obs",
        "match_alternativos",
    ]
    cols = [c for c in preferred if c in df.columns] + [
        c for c in df.columns if c not in preferred
    ]
    df = df[cols]

    xlsx = OUT / "cruce-informes-demolicion-habitable-2026-08-20.xlsx"
    xlsx_alt = OUT / "cruce-informes-demolicion-habitable-2026-08-20-v2.xlsx"
    csv_out = OUT / "cruce-informes-demolicion-habitable-2026-08-20.csv"
    df.to_csv(csv_out, index=False, encoding="utf-8-sig")

    resumen = pd.DataFrame(
        [
            {"metrica": "Informes PDF", "valor": len(df)},
            {
                "metrica": "Con texto usable",
                "valor": int((df["chars_texto_extraido"] > 200).sum())
                if "chars_texto_extraido" in df
                else "",
            },
            {
                "metrica": "Recomienda demolición (detectado)",
                "valor": int((df["recomienda_demolicion"] == "Sí").sum())
                if "recomienda_demolicion" in df
                else "",
            },
            {
                "metrica": "Match alto con Habitable",
                "valor": int(df["match_calidad"].astype(str).str.startswith("Alto").sum()),
            },
            {
                "metrica": "Match medio / a revisar",
                "valor": int(df["match_calidad"].astype(str).str.contains("Medio|Bajo|Sin", na=False).sum()),
            },
            {
                "metrica": "Match bajo / sin match",
                "valor": int(df["match_calidad"].isin(["Bajo — revisar", "Sin match"]).sum()),
            },
            {"metrica": "Habitable total filas", "valor": len(hab)},
            {"metrica": "Habitable etiqueta ROJO", "valor": len(hab_rojo)},
            {"metrica": "Corte Habitable", "valor": "2026-08-20 12:02:48"},
        ]
    )

    catalogo = pd.DataFrame(
        [
            {"nivel": "A", "campo": "habitable_id / certificado", "obligatorio": "Sí", "origen": "Habitable"},
            {"nivel": "A", "campo": "nombre_edificacion + alias", "obligatorio": "Sí", "origen": "Campo / informe"},
            {"nivel": "A", "campo": "direccion / municipio / parroquia", "obligatorio": "Sí", "origen": "Habitable + campo"},
            {"nivel": "A", "campo": "lat / lng GPS control", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "A", "campo": "fecha_inspeccion_detallada", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "A", "campo": "evaluadores / supervisor / ente", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "A", "campo": "etiqueta_fase1 + fecha_fase1", "obligatorio": "Sí", "origen": "Habitable"},
            {"nivel": "B", "campo": "uso / año construcción / n_torres", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "B", "campo": "num_pisos / sotanos / area_planta_m2", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "B", "campo": "sistema_estructural", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "B", "campo": "ocupacion_actual", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "B", "campo": "colindancias_criticas", "obligatorio": "Recomendado", "origen": "Campo"},
            {"nivel": "C", "campo": "piso_critico (nivel)", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "C", "campo": "% columnas grado III/IV/V", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "C", "campo": "dano_vigas_losas (flags + texto)", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "C", "campo": "inclinacion / verticalidad", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "C", "campo": "riesgo_no_estructural A/B/C", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "C", "campo": "peligro_aledanos", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "A0", "campo": "validacion_precarga (Sí/No/Parcial)", "obligatorio": "Sí", "origen": "Visita 2"},
            {"nivel": "A0", "campo": "score / banda / puesto ranking", "obligatorio": "Sí", "origen": "Sistema"},
            {"nivel": "D", "campo": "decision_D (D1–D5)", "obligatorio": "Sí", "origen": "Dictamen"},
            {"nivel": "D", "campo": "magnitud_M (M1–M4 o N/A)", "obligatorio": "Sí si D2", "origen": "Dictamen"},
            {"nivel": "D", "campo": "confirmacion_etiqueta_roja", "obligatorio": "Sí", "origen": "Dictamen"},
            {"nivel": "D", "campo": "prioridad + medidas_inmediatas", "obligatorio": "Sí", "origen": "Dictamen"},
            {"nivel": "D", "campo": "resumen_ejecutivo", "obligatorio": "Sí", "origen": "Dictamen"},
            {"nivel": "E", "campo": "fotos georref. / descritas (mín. 6)", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "E", "campo": "firmas elaboró/revisó/aprobó", "obligatorio": "Sí", "origen": "Campo"},
            {"nivel": "F", "campo": "estado_2da_ronda / versión_formato", "obligatorio": "Sí", "origen": "Sistema"},
        ]
    )

    ranking = build_ranking_rojo(hab_rojo, ids_en_informes)
    ranking_lg = ranking[
        ranking["estado"].astype(str).str.contains("Guaira|Vargas", case=False, na=False)
        | ranking["municipio"].astype(str).str.contains("Vargas", case=False, na=False)
    ].copy()
    ranking_lg = ranking_lg.reset_index(drop=True)
    ranking_lg["rank_gravedad_LaGuaira"] = ranking_lg.index + 1
    # mover columna rank local al frente
    cols_lg = ["rank_gravedad_LaGuaira"] + [
        c for c in ranking_lg.columns if c != "rank_gravedad_LaGuaira"
    ]
    ranking_lg = ranking_lg[cols_lg]

    indice = build_indice_pestanas()

    target_xlsx = xlsx
    try:
        # probar bloqueo
        with open(xlsx, "a"):
            pass
    except PermissionError:
        target_xlsx = xlsx_alt
        print(f"AVISO: {xlsx.name} está abierto; se guarda en {xlsx_alt.name}")

    with pd.ExcelWriter(target_xlsx, engine="openpyxl") as w:
        indice.to_excel(w, sheet_name="Indice_pestanas", index=False)
        df.to_excel(w, sheet_name="Cruce_informes", index=False)
        resumen.to_excel(w, sheet_name="Resumen", index=False)
        catalogo.to_excel(w, sheet_name="Catalogo_campos_propuestos", index=False)
        hab_rojo_lg.head(500).to_excel(
            w, sheet_name="Habitable_ROJO_LaGuaira_muestra", index=False
        )
        ranking.to_excel(w, sheet_name="Ranking_ROJO_nacional", index=False)
        ranking_lg.to_excel(w, sheet_name="Ranking_ROJO_LaGuaira", index=False)
        ranking.head(200).to_excel(w, sheet_name="Ranking_ROJO_Top200", index=False)

    # Control + vaciado digital (listas desplegables)
    import sys

    sys.path.insert(0, str(OUT))
    from excel_control_vaciado import build_control_2da_ronda, enrich_workbook

    control_df = build_control_2da_ronda(ranking_lg, df)
    enrich_workbook(target_xlsx, control_df)
    print("Control_2da_ronda filas:", len(control_df))

    print("Excel:", target_xlsx)
    print(
        "Ranking nacional:",
        len(ranking),
        "LaGuaira:",
        len(ranking_lg),
        "score>=60:",
        int((ranking["score_gravedad"] >= 60).sum()),
    )

    fmt = OUT / "Formato propuesto inspeccion detallada verificacion ROJO.docx"
    sis = OUT / "Esquema sistema vaciado inspecciones verificacion ROJO.docx"
    crit = OUT / "Criterios score gravedad etiquetas ROJAS.docx"
    build_formato_inspeccion_docx(fmt)
    build_esquema_sistema_docx(sis)
    build_criterios_score_docx(crit)
    print("Word formato:", fmt)
    print("Word esquema:", sis)
    print("Word criterios score:", crit)

    (OUT / "README-entregables-demolicion-ronda2.md").write_text(
        f"""# Entregables — segunda ronda verificación ROJO / demolición

Fecha: {date.today().isoformat()}

## Fuentes
- Informes PDF: carpeta Edificios Demolicion (Downloads)
- Habitable: export 2026-08-20_12:02:48

## Archivos
- `cruce-informes-demolicion-habitable-2026-08-20.xlsx` — listado operativo: cruce PDF, rankings, Control_2da_ronda
- `Ejemplo vaciado digital inspeccion 2da ronda Franco Mar.xlsx` — vaciado tipo informe (ejemplo + plantilla); acompaña el Word
- `Formato propuesto inspeccion detallada verificacion ROJO.docx`
- `Esquema sistema vaciado inspecciones verificacion ROJO.docx`
- `Criterios score gravedad etiquetas ROJAS.docx`

## Nota
Los match Habitable son heurísticos (nombre/GPS); revisar `match_calidad` y `match_alternativos`.
El score de gravedad prioriza cola de trabajo; no es dictamen de demolición.
El Excel de ejemplo de vaciado va junto al Word de formato; el Excel de cruce/ranking es el de control de casos.
""",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
