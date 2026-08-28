# -*- coding: utf-8 -*-
"""Genera 3 Word UTF-8: comparación Funvisis vs propuesta; informe consolidado; guía sistema juniors."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIRS = [
    Path(r"F:\servidor\2da ronda"),
    Path(
        r"C:\Users\Angel\Projects\clients\comision-presidencial-habitabilidad"
        r"\bi-habitable\docs\demolicion-ronda2"
    ),
]


def style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def p(doc: Document, text: str, bold: bool = False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return para


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = htxt
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), header_fill)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def save_all(doc: Document, filename: str) -> None:
    for d in OUT_DIRS:
        d.mkdir(parents=True, exist_ok=True)
        path = d / filename
        doc.save(str(path))
        print("OK", path)


def title_block(doc: Document, line1: str, line2: str) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(line1)
    r.bold = True
    r.font.size = Pt(12)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(line2)
    r2.bold = True
    r2.font.size = Pt(14)


def build_comparacion() -> None:
    doc = Document()
    style_doc(doc)
    title_block(
        doc,
        "COMISIÓN PRESIDENCIAL PARA LA EVALUACIÓN DE HABITABILIDAD",
        "Comparación técnica: propuesta 2.ª ronda (20/08) vs Guía Funvisis WRAP (25/08)",
    )
    p(
        doc,
        "Documento de trabajo · 27/08/2026 · Etapa inicial del sistema de seguimiento y vaciado Fase II",
    )
    p(
        doc,
        "Fuentes: Formato propuesto inspección detallada verificación ROJO; Criterios score gravedad; "
        "Ejemplo vaciado Franco Mar; Excel cruce operativo; Guía rápida llenado informe WRAP "
        "(Funvisis, 25/08/2026); Protocolo técnico evaluación/rehabilitación rev.6; "
        "Lineamientos apuntalamiento/estructuras/mampostería Sismo24J.",
    )

    doc.add_heading("1. Veredicto ejecutivo", level=1)
    p(
        doc,
        "Sí se alinean en el propósito de fondo, pero no son el mismo tipo de documento. "
        "Funvisis aporta la guía técnica normativa del informe de campo (cómo describir daño y qué "
        "reparación sugerir). La propuesta del 20/08 aporta el puente operativo hacia un sistema: "
        "precarga Habitable, ranking de cola, decisión de control D1–D5, magnitud M1–M4 y vaciado "
        "digital comparable.",
    )
    bullets(
        doc,
        [
            "Alineación alta en: Fase II post-etiqueta; estandarizar informes heterogéneos; "
            "daño estructural + mampostería; fotos de evidencia; identificación del inmueble; "
            "referencia a lineamientos CPEH/Funvisis.",
            "Diferencia principal: Funvisis = cómo escribir el diagnóstico técnico y elegir "
            "procedimientos VIG/COL/MAM. Propuesta 20/08 = cómo gobernar la cola caso a caso "
            "(Habitable → visita 2 → dictamen de control).",
            "Conclusión de diseño: Funvisis es la guía fundamental del cuerpo técnico del informe. "
            "La propuesta 20/08 es la capa de proceso/sistema que Funvisis no cubre (y no pretende cubrir).",
        ],
    )

    doc.add_heading("2. Qué propone cada lado", level=1)
    doc.add_heading("2.1 Propuesta inicial (20/08)", level=2)
    bullets(
        doc,
        [
            "Problema: muchos informes buenos pero distintos → no se puede listar cuántos van a "
            "demolición, cuántos a reparación y qué tan grande es esa reparación.",
            "Puente con Fase 1 Habitable: precarga (ID, certificado, etiqueta, riesgos A/B/C, acciones) "
            "+ score 0–100 / bandas de prioridad.",
            "Estructura A0–E: precarga+ranking → identidad → ficha edificio → matriz daño + texto libre "
            "→ decisión D1–D5 (+ M1–M4 si repara) → fotos/firmas + resumen ejecutivo.",
            "Excel operativo: Ranking_ROJO + Control_2da_ronda; Excel didáctico Franco Mar.",
            "Enfoque: control de proceso y datos estructurados para tablero/sistema.",
        ],
    )
    doc.add_heading("2.2 Funvisis — Guía WRAP (25/08)", level=2)
    bullets(
        doc,
        [
            "Problema: ambigüedad e inconsistencia en descripción técnica de fallas → retrasos y reinspecciones.",
            "Alcance explícito: Fase II detallada sobre edificios ya AMARILLO o ROJO; no sustituye "
            "inspección rápida ni etiquetas preliminares.",
            "Cuerpo del informe: datos del inmueble; descripción estructural (uso, año, niveles, "
            "geometría, irregularidades, dimensiones vigas/columnas/muros); clasificación daño "
            "estructural por mecanismo (vigas/columnas/muros) + nivel leve/moderado/severo; daño "
            "mampostería; diagnóstico narrativo; daños preexistentes; escaleras/evacuación.",
            "Núcleo diferencial: checklists de rehabilitación/reparación con procedimientos "
            "VIG-01..04, COL-01..05, MAM-01..04 ligados a lineamientos técnicos.",
            "Cierre: consideraciones (ensayos, modelaje no lineal, reforzamiento); semáforo; "
            "posibles salidas tipificadas (evacuar / reforzamiento global / demolición estructural).",
            "Enfoque: homogeneizar ingeniería de campo y vincular hallazgo → procedimiento técnico.",
        ],
    )

    doc.add_heading("3. Matriz de alineación", level=1)
    add_table(
        doc,
        ["Dimensión", "Propuesta 20/08", "Funvisis WRAP", "¿Alineado?"],
        [
            ["Fase II post semáforo", "Sí (énfasis ROJO)", "Sí (AMARILLO y ROJO)", "Sí (Funvisis más amplio)"],
            ["Estandarizar informes", "Sí (molde + casillas)", "Sí (guía + matrices)", "Sí"],
            ["Identidad inmueble", "A + precarga Habitable", "Tablas datos + uso + contactos", "Sí"],
            [
                "Geometría / sistema",
                "Parte B resumida",
                "Muy detallada (ejes, dims.)",
                "Parcial (Funvisis más rico)",
            ],
            [
                "Daño estructural",
                "Matriz % columnas, piso crítico, inclinación + texto",
                "Mecanismo + nivel por elemento + fotos",
                "Parcial → complementar",
            ],
            [
                "Mampostería",
                "Mencionada / riesgos no estructurales",
                "Mecanismo + nivel + MAM-01..04",
                "Funvisis más completo",
            ],
            [
                "Reparación sugerida",
                "Magnitud M1–M4 (tamaño)",
                "Procedimientos VIG/COL/MAM concretos",
                "Complementarios",
            ],
            [
                "Decisión demoler",
                "Código D1–D5 explícito",
                "Señales (evacuar / refuerzo / demolición)",
                "Parcial → unificar códigos",
            ],
            [
                "Puente Habitable Fase 1",
                "Precarga + validar/corregir",
                "No modela importación de planilla",
                "Solo propuesta 20/08",
            ],
            [
                "Cola / prioridad",
                "Score 0–100 + Ranking_ROJO",
                "No es ranking operativo",
                "Solo propuesta 20/08",
            ],
            [
                "Vaciado digital / sistema",
                "Excel + niveles A–F",
                "Documento Word de campo",
                "Solo propuesta 20/08",
            ],
            ["Fotos", "Mínimo 6 tipificadas", "Obligatorias por mecanismo", "Sí"],
            ["Acero estructural", "Posible en sistema", "Fuera de alcance (esta versión)", "Anotar límite"],
        ],
    )

    doc.add_heading("4. Diferencias que importan para el sistema", level=1)
    bullets(
        doc,
        [
            "Funvisis no define el flujo Habitable → cola → ficha → aprobación. Sin esa capa no hay "
            "sistema de seguimiento.",
            "La propuesta 20/08 no detalla mecanismos de falla ni checklists VIG/COL/MAM. Sin Funvisis "
            "el informe queda débil técnicamente frente a lineamientos oficiales.",
            "Funvisis incluye AMARILLO; la propuesta inicial se centró en ROJO/demolición. El sistema "
            "inicial puede nacer en ROJO y abrir AMARILLO en una ola 2 de producto.",
            "M1–M4 mide tamaño de obra para gobernanza; VIG/COL/MAM mide procedimiento técnico. Ambos "
            "deben coexistir: primero diagnóstico Funvisis, luego decisión D + magnitud M.",
            "El score de gravedad es herramienta de priorización, no dictamen. Funvisis tampoco lo "
            "contradice: simplemente no lo usa.",
        ],
    )

    doc.add_heading("5. Recomendación de base para construir el sistema", level=1)
    bullets(
        doc,
        [
            "Guía fundamental del informe de campo = Funvisis WRAP + lineamientos CPEH (concreto, "
            "mampostería, apuntalamiento) + Protocolo fases 1–4.",
            "Capa de proceso/datos = propuesta 20/08 (precarga Habitable, ranking, D1–D5, M1–M4, "
            "estados de flujo, resumen ejecutivo).",
            "Puente de producto mínimo: un formulario digital que (1) abre caso desde Ranking_ROJO, "
            "(2) obliga secciones Funvisis de daño y procedimientos, (3) cierra con decisión D/M y firmas.",
            "No intentar un «sistema de proyectos de reconstrucción» completo en la v1: el MVP es "
            "seguimiento + vaciado + dictamen de control + exportes.",
        ],
    )

    doc.add_heading("6. Riesgos si se elige solo un lado", level=1)
    add_table(
        doc,
        ["Si solo…", "Riesgo"],
        [
            [
                "Solo Funvisis (Word)",
                "Vuelve la heterogeneidad digital; no hay cola ni KPIs demoler/reparar; "
                "difícil auditar miles de ROJO.",
            ],
            [
                "Solo propuesta 20/08",
                "Dictámenes comparables pero débiles frente a lineamientos; inspectores sin "
                "checklist de reparación homologado.",
            ],
            [
                "Fusionar sin jerarquía",
                "Formulario interminable; abandono de campo. Hay que priorizar campos "
                "obligatorios vs opcionales.",
            ],
        ],
    )
    p(
        doc,
        "Siguiente entregable: Prototipo de informe consolidado Fase II (Funvisis como guía + capa de control 20/08).",
        bold=True,
    )
    save_all(doc, "2026-08-27 Comparacion Funvisis WRAP vs propuesta 2da ronda.docx")


def build_informe() -> None:
    doc = Document()
    style_doc(doc)
    title_block(
        doc,
        "COMISIÓN PRESIDENCIAL PARA LA EVALUACIÓN DE HABITABILIDAD",
        "PROTOTIPO DE INFORME CONSOLIDADO — INSPECCIÓN DETALLADA FASE II",
    )
    p(
        doc,
        "Versión 0.1 · 27/08/2026 · Guía fundamental: Funvisis WRAP 25/08 · Capa de proceso: "
        "propuesta 20/08 · Referencia de campo: informes tipo Franco Mar / Capri / OPPPE",
    )
    p(
        doc,
        "Este documento es el molde de captura (papel/PDF/sistema). No sustituye el proyecto "
        "ejecutivo de demolición ni el diseño de refuerzo.",
    )

    doc.add_heading("0. Principios del prototipo", level=1)
    bullets(
        doc,
        [
            "Funvisis define el lenguaje técnico del daño y las alternativas de reparación (VIG/COL/MAM).",
            "La capa Habitable/ranking define de dónde viene el caso y con qué prioridad se atiende.",
            "La decisión de control (D1–D5 + M1–M4) cierra el caso para mesas de trabajo y tableros.",
            "Campos estructurados = comparabilidad. Texto libre = juicio del ingeniero. Fotos = evidencia mínima.",
            "MVP: ROJO prioritario. Extensión posterior: AMARILLO (ya previsto por Funvisis).",
        ],
    )

    doc.add_heading("1. Portada institucional", level=1)
    bullets(
        doc,
        [
            "Título: Inspección detallada Fase II — Edificio «NOMBRE» — Eventos sísmicos 24/06/2026.",
            "Equipo inspector: nombre, C.I., CIV, ente; fecha/hora inicio y fin.",
            "Dictamen Habitable Fase 1 (precargado) y dictamen Fase II (resultado de esta visita).",
        ],
    )

    doc.add_heading("2. Bloque A0 — Precarga Habitable + ranking (solo lectura + validación)", level=1)
    p(doc, "Obligatorio en sistema digital. En papel, se pega la ficha previa.")
    add_table(
        doc,
        ["Campo", "Origen", "Acción inspector"],
        [
            ["ID / certificado Habitable", "Export Fase 1", "Confirmar edificio correcto"],
            ["Nombre, dirección, GPS, municipio/parroquia", "Fase 1", "Validar / corregir"],
            ["Etiqueta, riesgos A/B/C, piso crítico, acciones", "Planilla ANIH", "Confirmar o corregir"],
            ["Score 0–100, banda, puesto, detalle", "Ranking ROJO", "Usar como prioridad; no como dictamen"],
            ["Validación A0.3", "Visita 2", "Correcto / Corregir / Evidencia insuficiente"],
        ],
    )

    doc.add_heading("3. Bloque A — Datos del inmueble (Funvisis)", level=1)
    bullets(
        doc,
        [
            "Dirección, coordenadas UTM (19P) o GPS control, contacto, ocupación (n.º personas).",
            "Uso (vivienda uni/multi, comercial, educativo, salud, etc. — tabla Funvisis).",
            "Persona contacto y teléfono.",
        ],
    )

    doc.add_heading("4. Bloque B — Descripción estructural (Funvisis, simplificada para MVP)", level=1)
    bullets(
        doc,
        [
            "Obligatorios MVP: sistema constructivo (tabla Funvisis), año/rango construcción, "
            "N.º pisos / semisótano / sótano, N.º aptos/locales, área planta tipo y área total "
            "(si se conoce), geometría planta/elevación, irregularidades clave (entrepiso débil, "
            "columnas cortas, adosamiento…).",
            "Recomendados (Fase II completa): conteo vigas/columnas/muros en planta tipo; "
            "dimensiones por ejes (tablas Funvisis).",
            "En MVP de campo se permiten croquis + «dimensiones representativas» si no hay plano; "
            "el sistema marca calidad_dato = estimado.",
        ],
    )

    doc.add_heading("5. Bloque C — Clasificación y diagnóstico del daño estructural (Funvisis = guía)", level=1)
    p(doc, "Por tipología de elemento (vigas / columnas / muros-pantallas):")
    bullets(
        doc,
        [
            "Mecanismo de falla (flexión, corte, corte+flexión, compresión, adherencia, etc.).",
            "Nivel de daño: Sin daño o leve / Moderado / Severo (umbrales de grieta Funvisis).",
            "Ubicación (nivel/eje) + al menos 2 fotos (vista completa + detalle).",
            "Diagnóstico narrativo del daño estructural de la edificación (texto libre obligatorio).",
            "Diagnóstico de escaleras / vías de evacuación.",
            "Daños preexistentes o anteriores al 24/06/2026 (separar de daño sísmico).",
        ],
    )
    p(
        doc,
        "Campos puente (de la propuesta 20/08, útiles para tablero): piso crítico; % columnas en "
        "daño grave; inclinación (sí/no + Δ); peligro aledaños; riesgos no estructurales A/B/C.",
    )

    doc.add_heading("6. Bloque D — Daño en mampostería (Funvisis)", level=1)
    bullets(
        doc,
        [
            "Mecanismo (tracción/flexión, corte, fuera de plano, compresión).",
            "Nivel leve / moderado / severo (umbrales Funvisis).",
            "Fotos fachada/pared + detalle; diagnóstico narrativo.",
        ],
    )

    doc.add_heading("7. Bloque E — Métodos de rehabilitación/reparación sugeridos (Funvisis)", level=1)
    p(doc, "Checklists con X según hallazgo. El sistema guarda códigos de procedimiento seleccionados.")
    add_table(
        doc,
        ["Familia", "Códigos", "Idea"],
        [
            ["Vigas", "VIG-01 … VIG-04", "Inyección → reparación local → rótula/núcleo"],
            [
                "Columnas",
                "COL-01 … COL-05 (+ crítica)",
                "Sellado → inyección → recubrimiento → reconstrucción; fuera de alcance",
            ],
            [
                "Mampostería",
                "MAM-01 … MAM-04",
                "Fisuras → inyección → llaveado → tabiquería liviana en evacuación",
            ],
        ],
    )
    p(
        doc,
        "Nota: seleccionar procedimiento no implica que la obra esté presupuestada; es "
        "recomendación técnica alineada a lineamientos.",
    )

    doc.add_heading("8. Bloque F — Decisión de control (capa 20/08, obligatoria para el sistema)", level=1)
    add_table(
        doc,
        ["Código", "Decisión", "Siguiente paso típico"],
        [
            ["D1", "Demoler", "Proyecto demolición controlada + perímetro"],
            ["D2", "Reparar / reconstruir", "Obligatorio M1–M4 + procedimientos VIG/COL/MAM"],
            ["D3", "Más estudios", "Ensayos / modelaje / reinspección"],
            ["D4", "Escombros / ya colapsó", "Remoción y control de sitio"],
            ["D5", "Inhabitabilidad + vigilancia", "Perímetro y monitoreo; sin demolición aún"],
        ],
    )
    add_table(
        doc,
        ["Magnitud (si D2)", "Definición operativa"],
        [
            ["M1", "Reparación pequeña / local"],
            ["M2", "Reparación importante (varios elementos o un nivel; apuntalamiento)"],
            ["M3", "Reconstrucción parcial (p. ej. piso crítico)"],
            ["M4", "Reconstrucción / refuerzo mayor"],
        ],
    )
    bullets(
        doc,
        [
            "Prioridad operativa: Inmediata / Alta / Programable.",
            "Medidas inmediatas: acordonar, excluir, apuntalar temporal, monitorear vecinos, etc.",
            "Justificación libre obligatoria (por qué D/M; contraste con ranking si difiere).",
            "Puente Funvisis↔D: si hay «condición crítica / fuera de alcance» en columnas o "
            "demolición estructural marcada → el sistema sugiere D1 (el ingeniero confirma).",
        ],
    )

    doc.add_heading(
        "9. Bloque G — Estimación referencial de reconstrucción (opcional en MVP, diseñada desde ya)",
        level=1,
    )
    p(
        doc,
        "No es presupuesto oficial. Es orden de magnitud para mesa técnica, usando área (m²), "
        "magnitud M, tipología y factores editables (USD/m² por escenario). Se presenta siempre "
        "como estimación preliminar.",
    )
    bullets(
        doc,
        [
            "Entradas: área total o afectada; M1–M4 o D1/D4; % planta afectada; requiere refuerzo global (sí/no).",
            "Salida: rango bajo–alto + supuestos. Editable por coordinador técnico.",
        ],
    )

    doc.add_heading("10. Bloque H — Evidencia, firmas y resumen ejecutivo", level=1)
    bullets(
        doc,
        [
            "Fotos mínimas: vista general; piso crítico; columnas; vigas/losas; mampostería; "
            "aledaños/evacuación; etiqueta.",
            "Croquis recomendado.",
            "Firmas: elaboró / revisó / aprobó (CIV).",
            "Resumen ejecutivo 8–12 líneas: Fase 1 + ranking + validación + hallazgo + D/M + próxima acción.",
        ],
    )

    doc.add_heading("11. Mapa de secciones → sistema digital", level=1)
    add_table(
        doc,
        ["Sección informe", "Módulo sistema", "Obligatorio MVP"],
        [
            ["A0 Precarga", "Caso + import Habitable", "Sí"],
            ["A–B Datos/estructura", "Ficha edificio", "Sí (subset)"],
            ["C–D Daño", "Hallazgos + fotos", "Sí"],
            ["E Procedimientos", "Checklist VIG/COL/MAM", "Sí (al menos 1 familia tocada)"],
            ["F Decisión D/M", "Dictamen + workflow", "Sí"],
            ["G Estimación USD", "Calculadora referencial", "No (piloto)"],
            ["H Firmas/PDF", "Export PDF + estados", "Sí"],
        ],
    )

    doc.add_heading("12. Ejemplo de cierre (Franco Mar — didáctico)", level=1)
    p(
        doc,
        "Precarga: ROJO 20/07/2026, score 51 (Media). Visita detallada: piso crítico Nivel 1, "
        ">50% columnas graves, pérdida de verticalidad. Procedimientos: condición crítica / "
        "fuera de alcance en columnas. Decisión: D1 Demoler, prioridad inmediata. Aprendizaje: "
        "el ranking prioriza; Funvisis describe; D1 cierra el control.",
    )
    p(
        doc,
        "Documento hermano: propuesta de creación del sistema (equipo junior) — mismo paquete 27/08/2026.",
        bold=True,
    )
    save_all(doc, "2026-08-27 Prototipo informe consolidado Fase II.docx")


def build_sistema() -> None:
    doc = Document()
    style_doc(doc)
    title_block(
        doc,
        "PROPUESTA DE SISTEMA — SEGUIMIENTO Y VACIADO FASE II",
        "Guía de trabajo para equipo de desarrollo junior · Marco Django · 27/08/2026",
    )
    p(
        doc,
        "Audiencia: programadores junior. Objetivo: entender por qué Django, qué se construye en "
        "la v1, y cómo modelar el informe consolidado (Funvisis + capa Habitable/D-M).",
    )

    doc.add_heading("1. El problema en lenguaje simple", level=1)
    bullets(
        doc,
        [
            "Ya existe Fase 1 (Habitable): miles de edificios con etiqueta (verde/amarillo/rojo).",
            "Ahora hay Fase 2: visitar en detalle (empezando por ROJO), validar lo anterior, "
            "describir daños con criterio Funvisis y decidir demoler / reparar / estudiar / "
            "escombros / vigilar.",
            "Hoy eso vive en PDF/Word/Excel distintos. No hay usuarios, estados ni base de datos única.",
            "No pedimos un ERP de reconstrucción nacional. Pedimos un puente: cola → ficha → dictamen → exportes.",
        ],
    )

    doc.add_heading("2. Qué NO es este proyecto (para no sobre-diseñar)", level=1)
    bullets(
        doc,
        [
            "No es BIM ni modelación estructural no lineal.",
            "No es contabilidad de obras ni licitaciones.",
            "No reemplaza Habitable 1×10 ni el tablero PDNA Streamlit (pueden integrarse después).",
            "No sustituye los lineamientos técnicos PDF: el sistema los referencia y captura checklists.",
        ],
    )

    doc.add_heading("3. Selección del marco de trabajo", level=1)
    doc.add_heading(
        "3.1 Recomendación: Django + PostgreSQL + (opcional) HTMX/Alpine o templates Django",
        level=2,
    )
    p(doc, "Por qué Django encaja aquí:")
    bullets(
        doc,
        [
            "Usuarios e inspectores: Django Auth + grupos (Inspector, Supervisor, Coordinador, Solo lectura).",
            "Formularios largos y validaciones: ModelForm / forms por secciones del informe.",
            "Admin inmediato para operación (Django Admin) mientras se pulen pantallas.",
            "ORM maduro para el modelo Caso / Visita / Hallazgo / Foto / Dictamen.",
            "Ecosistema conocido en el programa CPEH (ya hay experiencia Django en el ecosistema web).",
            "Export PDF (WeasyPrint/xhtml2pdf) y Excel (openpyxl) sin inventar stack.",
            "Despliegue predecible (Render/Docker) similar a otros productos del ecosistema.",
        ],
    )
    doc.add_heading("3.2 Alternativas consideradas (y por qué no son el MVP)", level=2)
    add_table(
        doc,
        ["Opción", "Pros", "Contras para este MVP"],
        [
            ["Solo Excel + Drive", "Rápido", "Sin usuarios reales, sin auditoría, rompe a escala"],
            [
                "Streamlit",
                "Rápido para BI",
                "Malo para multi-usuario, permisos y formularios de campo",
            ],
            ["Reflex / SPA Python", "UI moderna", "Más curva; overkill para juniors en v1"],
            [
                "Firebase + app móvil",
                "Offline fácil",
                "Modelo de datos y reportes institucionales más difíciles",
            ],
            [
                "Laravel / Rails",
                "También válidos",
                "Menos alineado al stack actual del equipo Python",
            ],
        ],
    )
    doc.add_heading("3.3 Arquitectura lógica (una frase)", level=2)
    p(
        doc,
        "Habitable (origen) → import/ranking → Caso en cola → Visita Fase II (formulario por "
        "bloques Funvisis) → Dictamen D/M → PDF/Excel → tablero de avance.",
    )

    doc.add_heading("4. Roles y permisos (v1)", level=1)
    add_table(
        doc,
        ["Rol", "Puede"],
        [
            ["Inspector", "Tomar casos asignados; editar borrador; subir fotos; enviar a revisión"],
            ["Supervisor", "Revisar, devolver, aprobar dictamen"],
            ["Coordinador", "Asignar cola, ver KPIs, editar factores de estimación"],
            ["Admin", "Usuarios, catálogos VIG/COL/MAM, import Habitable"],
        ],
    )

    doc.add_heading("5. Modelo de datos (explicado para juniors)", level=1)
    p(
        doc,
        "Piensen en tablas. Cada fila es un registro. Relaciones: un Caso tiene muchas Visitas; "
        "una Visita tiene muchos Hallazgos y Fotos; una Visita tiene un Dictamen.",
    )
    add_table(
        doc,
        ["Modelo", "Campos clave", "Notas"],
        [
            [
                "BuildingCase",
                "habitable_id, certificado, nombre, gps, etiqueta_f1, score, banda, estado_cola",
                "Puente con Fase 1",
            ],
            [
                "Phase1Snapshot",
                "JSON o columnas de riesgos A/B/C, acciones, piso_critico",
                "Inmutable histórico",
            ],
            [
                "InspectionVisit",
                "caso FK, fecha, equipo, estado_workflow",
                "Borrador→Revisión→Aprobado",
            ],
            [
                "StructureProfile",
                "sistema, pisos, sótanos, uso, irregularidades…",
                "Bloque B Funvisis",
            ],
            [
                "StructuralFinding",
                "elemento, mecanismo, nivel_dano, ubicacion, notas",
                "Bloque C",
            ],
            ["MasonryFinding", "mecanismo, nivel, notas", "Bloque D"],
            [
                "RepairRecommendation",
                "codigo (VIG-01…), seleccionado bool, notas",
                "Bloque E",
            ],
            [
                "ControlDecision",
                "codigo_D, magnitud_M, prioridad, medidas M2M, justificacion",
                "Bloque F",
            ],
            ["CostEstimate", "rango_low, rango_high, supuestos", "Opcional MVP"],
            ["PhotoEvidence", "tipo, archivo, caption", "Bloque H"],
            ["AuditLog", "user, action, timestamp, diff", "Trazabilidad"],
        ],
    )
    p(
        doc,
        "Estados de workflow sugeridos: pendiente_asignacion → en_borrador → en_revision → "
        "aprobado → publicado_export. Un caso aprobado no se edita sin reapertura supervisada.",
    )

    doc.add_heading("6. Pantallas mínimas (MVP)", level=1)
    bullets(
        doc,
        [
            "Login.",
            "Cola / Ranking: filtros banda, municipio, asignado a mí, estado.",
            "Ficha caso: A0 precarga + validación.",
            "Wizard de visita: pestañas A–H según prototipo de informe (guardar parcial).",
            "Galería de fotos con tipos obligatorios.",
            "Cierre: D/M + enviar a revisión.",
            "Bandeja supervisor: aprobar / devolver con comentario.",
            "Export: PDF institucional + Excel fila plana para tablero.",
            "Admin: catálogo procedimientos Funvisis; import CSV/Parquet Habitable.",
        ],
    )

    doc.add_heading("7. Integraciones", level=1)
    bullets(
        doc,
        [
            "Entrada: export Habitable (CSV/Parquet) + score/ranking (mismo algoritmo documentado 20/08).",
            "Salida: Excel compatible con Control_2da_ronda; PDF del informe consolidado.",
            "Futuro: API hacia BI Habitable / PDNA; app móvil offline (Fase posterior).",
        ],
    )

    doc.add_heading("8. Plan de entregas (sprints cortos)", level=1)
    add_table(
        doc,
        ["Sprint", "Entrega", "Criterio de hecho"],
        [
            ["0 (1 sem)", "Repo Django, auth, modelos Caso/Visita, Admin", "Crear caso manual + login"],
            ["1", "Import Habitable + cola ranking", "Filtrar ROJO por banda Alta/Muy alta"],
            ["2", "Wizard A0–B + fotos", "Inspector completa ficha básica"],
            ["3", "Hallazgos C–D + checklists E Funvisis", "Guardar VIG/COL/MAM"],
            ["4", "Dictamen F + workflow supervisor + PDF", "Caso aprobado exporta PDF"],
            ["5", "Excel control + KPI demoler/reparar", "Dashboard simple"],
            ["6 (piloto)", "Estimación G + mejoras UX", "Piloto 20–50 edificios reales"],
        ],
    )

    doc.add_heading("9. Stack técnico concreto (v1)", level=1)
    bullets(
        doc,
        [
            "Python 3.11+, Django 5.x, PostgreSQL 15+, Pillow (fotos), openpyxl, WeasyPrint o similar PDF.",
            "Almacenamiento archivos: local/S3 compatible.",
            "Frontend: Django templates + Bootstrap 5 (mismo lenguaje visual que otros portales CPEH) "
            "+ HTMX opcional para guardar secciones sin recargar todo.",
            "Tests: pytest-django en modelos de decisión y permisos.",
            "Entornos: local Docker Compose; staging; prod detrás de HTTPS.",
        ],
    )

    doc.add_heading("10. Reglas de negocio que el código debe respetar", level=1)
    bullets(
        doc,
        [
            "No se crea visita Fase II sin habitable_id (o excepción supervisada documentada).",
            "Si decision=D2 → magnitud_M obligatoria y al menos un procedimiento de reparación seleccionado.",
            "Si hay hallazgo de columna «fuera de alcance» → UI sugiere D1 pero no fuerza sin confirmación.",
            "Score/ranking nunca escribe solo el dictamen D.",
            "Montos de estimación siempre con leyenda «preliminar / no presupuestario».",
            "Auditoría: quién cambió D/M y cuándo.",
        ],
    )

    doc.add_heading("11. Organización del repositorio (sugerida)", level=1)
    bullets(
        doc,
        [
            "apps/accounts — usuarios/roles",
            "apps/cases — BuildingCase, import, cola",
            "apps/inspections — visita, hallazgos, fotos, recomendaciones",
            "apps/decisions — ControlDecision, workflow",
            "apps/exports — PDF/Excel",
            "apps/catalogs — VIG/COL/MAM, usos, sistemas",
            "docs/ — prototipo informe + comparación Funvisis (este paquete)",
        ],
    )

    doc.add_heading("12. Criterios de calidad para juniors", level=1)
    bullets(
        doc,
        [
            "Cada PR: migraciones + test mínimo del modelo tocado.",
            "No hardcodear textos de procedimientos: viven en Catalog (para actualizar si Funvisis revisa la guía).",
            "Validar formularios en servidor (nunca confiar solo en el navegador).",
            "Fotos: límite de tamaño + tipos JPEG/PNG; nombres no sensibles en URL pública.",
            "Mensajes UI en español (es-VE).",
        ],
    )

    doc.add_heading("13. Definición de éxito del punto inicial", level=1)
    bullets(
        doc,
        [
            "Un inspector inicia sesión, toma un caso del Ranking_ROJO, completa el informe "
            "consolidado y obtiene PDF.",
            "Un supervisor aprueba y el caso sale de «pendiente verificación».",
            "Un coordinador ve conteos: D1/D2/D3/D4/D5 y M1–M4.",
            "Los datos de Fase 1 quedan trazados (qué se validó/corrigió).",
        ],
    )

    doc.add_heading("14. Próximos pasos de producto (fuera del código)", level=1)
    bullets(
        doc,
        [
            "Validar este prototipo de informe con Funvisis/comité técnico (1 reunión de homologación).",
            "Congelar catálogo VIG/COL/MAM v1 desde la Guía 25/08.",
            "Elegir piloto territorial (p. ej. subset La Guaira banda Alta).",
            "Decidir hosting y responsables de import Habitable semanal.",
        ],
    )
    p(
        doc,
        "Documentos base de este paquete: (1) Comparación Funvisis vs propuesta 20/08; "
        "(2) Prototipo informe consolidado Fase II; (3) esta guía de sistema.",
        bold=True,
    )
    save_all(doc, "2026-08-27 Propuesta sistema seguimiento Fase II - guia equipo junior.docx")


def main() -> None:
    build_comparacion()
    build_informe()
    build_sistema()
    # Verificación rápida de acentos
    sample = OUT_DIRS[0] / "2026-08-27 Comparacion Funvisis WRAP vs propuesta 2da ronda.docx"
    d = Document(str(sample))
    text = "\n".join(p.text for p in d.paragraphs[:8] if p.text.strip())
    assert "COMISIÓN" in text and "Comparación" in text and "Guía" in text, text[:200]
    print("VERIFICACION_OK acentos presentes")


if __name__ == "__main__":
    main()
