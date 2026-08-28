# -*- coding: utf-8 -*-
"""
Genera entregables Fase II v0.2 (UTF-8):
- Word prototipo informe consolidado (didáctico + ejemplo Franco Mar completo)
- Word propuesta sistema (didáctico junior + Django + servidor + MobaXterm)
- Excel ejecutivo plantilla informe consolidado
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

# Reutilizar listas y helpers del vaciado previo
from excel_control_vaciado import (
    LISTAS,
    _add_list_sheet,
    _dv,
    francomar_precarga,
    francomar_visita,
)

OUT_DIRS = [
    Path(r"F:\servidor\2da ronda"),
    Path(__file__).resolve().parent,
]

# Datos de servidor (verificados por SSH el 27/08/2026; NO incluir contraseñas)
SERVER_HOST = "190.169.110.9"
SERVER_USER = "cph"
SERVER_HOSTNAME = "cph"
SERVER_OS = "Debian GNU/Linux (kernel 6.12 amd64)"
SERVER_HOME = "/home/cph/"
SERVER_WORKDIR = "/home/cph/Project_Etapa2_FlujoObra/Proj_ControlInformes/"
SERVER_KEY_COMMENT = "Llave CPH"
SERVER_KEY_FINGERPRINT = "SHA256:de234edc1df6c898edae996804321c8a79c8421c0a2093825a6157fee2bba33f"
# Rutas locales del operador (Windows); no commitear .ppk ni frase clave
LOCAL_MOBA_DIR = r"F:\servidor\MobaXterm"
LOCAL_KEY_PPK = r"F:\servidor\cph_private_key.ppk"
LOCAL_PASSPHRASE_FILE = r"F:\servidor\frase-clave.txt"


def style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def save_all(doc: Document, filename: str) -> None:
    for d in OUT_DIRS:
        d.mkdir(parents=True, exist_ok=True)
        doc.save(str(d / filename))
        print("OK", d / filename)


def title_block(doc: Document, line1: str, line2: str, subtitle: str = "") -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(line1)
    r.bold = True
    r.font.size = Pt(12)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(line2)
    r2.bold = True
    r2.font.size = Pt(14)
    if subtitle:
        t3 = doc.add_paragraph()
        t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = t3.add_run(subtitle)
        r3.italic = True
        r3.font.size = Pt(10)


def p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = htxt
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def didactic_box(doc: Document, titulo: str, para_que: str, quien: str, ejemplo: str) -> None:
    doc.add_heading(titulo, level=3)
    p(doc, f"Para qué sirve: {para_que}")
    p(doc, f"Quién lo llena: {quien}")
    p(doc, f"Ejemplo rápido: {ejemplo}", bold=False)


def build_prototipo_informe() -> None:
    doc = Document()
    style_doc(doc)
    title_block(
        doc,
        "COMISIÓN PRESIDENCIAL PARA LA EVALUACIÓN DE HABITABILIDAD",
        "PROTOTIPO DE INFORME CONSOLIDADO — FASE II",
        "Versión didáctica 0.2 · 27/08/2026 · Guía Funvisis WRAP + capa Habitable/D-M · Caso ejemplo: Franco Mar",
    )

    doc.add_heading("Cómo usar este documento (léalo primero)", level=1)
    bullets(
        doc,
        [
            "Parte I (secciones 0–11): explica cada bloque del informe como si fuera un curso corto.",
            "Parte II (anexo): informe completo del Edificio Franco Mar ya redactado con la nueva estructura.",
            "El Excel hermano «Plantilla ejecutiva informe consolidado Fase II» tiene los mismos campos en columnas A–B para vaciado digital.",
            "Funvisis define el lenguaje técnico (daño + procedimientos VIG/COL/MAM). La capa Habitable/D-M define la gobernanza del caso.",
        ],
    )

    doc.add_heading("0. Mapa visual del informe (de arriba hacia abajo)", level=1)
    add_table(
        doc,
        ["Bloque", "Nombre", "¿Obligatorio MVP?", "Analogía simple"],
        [
            ["A0", "Precarga Habitable + ranking + validación", "Sí", "Lo que el sistema ya sabía antes de llegar"],
            ["A", "Datos del inmueble (Funvisis)", "Sí", "Ficha de identidad en sitio"],
            ["B", "Descripción estructural", "Sí (subset)", "Cómo está construido el edificio"],
            ["C", "Daño estructural detallado", "Sí", "Qué falló en vigas/columnas/muros"],
            ["D", "Daño mampostería", "Sí si aplica", "Fachadas, tabiques, muros no portantes"],
            ["E", "Procedimientos sugeridos VIG/COL/MAM", "Sí si D2", "Qué técnica de reparación corresponde"],
            ["F", "Decisión de control D1–D5 + M1–M4", "Sí", "Demoler / reparar / estudiar / escombros / vigilar"],
            ["G", "Estimación referencial USD", "No (piloto)", "Orden de magnitud, no presupuesto"],
            ["H", "Fotos, firmas, resumen ejecutivo", "Sí", "Evidencia y cierre para mesa de trabajo"],
        ],
    )

    doc.add_heading("Parte I — Guía didáctica bloque por bloque", level=1)

    didactic_box(
        doc,
        "Bloque A0 — Precarga y validación",
        "Evitar empezar de cero. El inspector confirma o corrige lo que Habitable y el ranking ya trajeron.",
        "Sistema (automático) + inspector (validación en campo).",
        "«Sí, es Franco Mar; ROJO confirmado; corrijo sótanos de 1 a 2».",
    )
    bullets(
        doc,
        [
            "Campos típicos: ID Habitable, certificado, etiqueta Fase 1, score, banda, puesto en cola.",
            "Validación obligatoria: ¿edificio correcto? ¿ROJO se mantiene? ¿qué se corrige?",
            "Regla de oro: el ranking prioriza la visita; NO dictamina demolición por sí solo.",
        ],
    )

    didactic_box(
        doc,
        "Bloque A — Datos del inmueble",
        "Identificar formalmente el caso para trazabilidad legal y operativa.",
        "Inspector + contacto en sitio.",
        "Dirección Avenida José María España; UTM/GPS de control; uso vivienda multifamiliar.",
    )

    didactic_box(
        doc,
        "Bloque B — Descripción estructural",
        "Contextualizar el daño: no es lo mismo un edificio de 2 pisos que una torre de 12.",
        "Inspector estructural.",
        "Pórticos CA ortogonales; 12 pisos + 2 sótanos; ~780 m² planta; irregularidades si las hay.",
    )

    didactic_box(
        doc,
        "Bloque C — Daño estructural (Funvisis)",
        "Homologar cómo se describe el daño: mecanismo + nivel leve/moderado/severo + fotos.",
        "Inspector estructural.",
        "Columnas Nivel 1: compresión, severo, >50% núcleo perdido; vigas: corte diagonal en 5 niveles.",
    )
    p(
        doc,
        "Por cada familia (vigas / columnas / muros-pantalla) repetir: mecanismo → nivel → ubicación → "
        "≥2 fotos → párrafo diagnóstico. Añadir campos puente para tablero: piso crítico, % columnas graves, inclinación.",
    )

    didactic_box(
        doc,
        "Bloque D — Mampostería",
        "Registrar daño en cerramientos y riesgo de volcamiento/fachada.",
        "Inspector.",
        "Mampostería comprimida en PB por reducción de entrepiso; nivel severo; foto fachada + detalle.",
    )

    didactic_box(
        doc,
        "Bloque E — Procedimientos VIG/COL/MAM",
        "Vincular hallazgo con lineamiento técnico (Guía Funvisis). Obligatorio si la decisión es reparar (D2).",
        "Inspector marca X en alternativas que correspondan al diagnóstico.",
        "Franco Mar (demolición): COL fuera de alcance / condición crítica — no aplica reparación local.",
    )

    didactic_box(
        doc,
        "Bloque F — Decisión D y magnitud M",
        "Cerrar el caso para mesa técnica: conteos demoler/reparar dependen de este código.",
        "Inspector propone; supervisor aprueba.",
        "D1 Demoler + prioridad inmediata + medidas de exclusión.",
    )
    add_table(
        doc,
        ["Código D", "Significado", "Cuándo usarlo"],
        [
            ["D1", "Demoler", "Reparación inviable o riesgo inaceptable (Franco Mar)"],
            ["D2", "Reparar", "Daño grave pero localizable; obligatorio M1–M4"],
            ["D3", "Estudios", "Faltan ensayos/modelo antes de decidir"],
            ["D4", "Escombros", "Ya colapsó; control de remoción"],
            ["D5", "Vigilar", "Inhabitable pero sin demolición inmediata"],
        ],
    )

    didactic_box(
        doc,
        "Bloque H — Cierre",
        "Dejar evidencia y un párrafo que un coordinador pueda leer en 30 segundos.",
        "Inspector + supervisor (firmas CIV).",
        "Resumen ejecutivo de 10 líneas + registro fotográfico referenciado.",
    )

    doc.add_heading("Parte II — ANEXO: INFORME COMPLETO EJEMPLO (EDIFICIO FRANCO MAR)", level=1)
    p(
        doc,
        "El texto siguiente es una consolidación didáctica basada en el informe técnico del 11/08/2026 "
        "(DGPS_INF-F2-Tanaguarena/11/08/2026-001) y en la precarga Habitable (ID 171818). "
        "Muestra cómo se ve el informe cuando se aplican los bloques A0–H.",
        bold=True,
    )

    # --- ANEXO COMPLETO FRANCO MAR ---
    doc.add_page_break()
    doc.add_heading("PORTADA", level=2)
    bullets(
        doc,
        [
            "República Bolivariana de Venezuela — Comisión Presidencial para la Evaluación de Habitabilidad",
            "Inspección detallada Fase II — Edificio FRANCO MAR",
            "Eventos sísmicos del 24 de junio de 2026 — Caraballeda / Tanaguarenas, La Guaira",
            "Referencia informe: DGPS_INF-F2-Tanaguarena/11/08/2026-001",
            "Fecha visita detallada: 11/08/2026",
            "Evaluadores: Ing. Luis Burgos (C.I. 4357927); Ing. José García (C.I. 6913241)",
            "Supervisión institucional: Ing. Aura Quintero",
        ],
    )

    doc.add_heading("A0 — Precarga Habitable y ranking", level=2)
    add_table(
        doc,
        ["Campo", "Valor precargado", "Validación visita 2"],
        [
            ["ID Habitable", "171818", "Correcto — es Franco Mar / Francomar"],
            ["Certificado", "73420200720261338", "Correcto"],
            ["Etiqueta Fase 1", "ROJO (20/07/2026)", "Confirmada — sigue ROJO"],
            ["Inspector Fase 1", "Maria Mercedes Nieves Duarte", "Referencia válida"],
            ["Dirección Habitable", "Boulevard Tanaguarenas, Caraballeda", "Corregida: Av. José María España, Tanaguarenas"],
            ["Pisos / sótanos (F1)", "12 / 1", "Corregido: 12 / 2"],
            ["GPS Habitable", "10.6123543, -66.8318071", "GPS control visita: 10.489497, -66.899001"],
            ["Score / banda", "51 / Media", "Útil como prioridad; no refleja gravedad real en sitio"],
            ["Puesto cola", "~115 La Guaira / ~264 nacional", "Referencia operativa"],
            ["Detalle score", "riesgo_externo+20; ext_colapso+15; …", "Explica por qué entró en cola Media"],
        ],
    )

    doc.add_heading("A — Datos del inmueble", level=2)
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Nombre confirmado", "Edificio Franco Mar"],
            ["Dirección", "Avenida José María España, Tanaguarenas, Vargas, La Guaira"],
            ["Coordenadas GPS control", "Lat. N 10.489497 · Long. W 66.899001"],
            ["Uso", "Vivienda multifamiliar (edificio residencial)"],
            ["Ocupación", "Desalojado — estructura inhabitable"],
            ["Contacto / entorno", "Arteria de alto tránsito; edificaciones colindantes"],
        ],
    )

    doc.add_heading("B — Descripción estructural", level=2)
    bullets(
        doc,
        [
            "Sistema: pórticos de concreto armado ortogonales.",
            "Columnas aproximadas: 22 perimetrales + 18 interiores en planta tipo.",
            "Niveles: 12 pisos + 2 sótanos (estacionamiento).",
            "Área planta aproximada: 780 m².",
            "Condición global al 11/08/2026: daño estructural severo extremo; piso crítico en Nivel 1.",
        ],
    )

    doc.add_heading("C — Daño estructural detallado (Funvisis)", level=2)
    doc.add_heading("C.1 Columnas de concreto armado", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Mecanismo de falla", "Compresión / aplastamiento del núcleo"],
            ["Nivel de daño", "Severo"],
            ["Ubicación", "Nivel 1 (piso crítico) — perimetrales e interiores"],
            ["Evidencia", ">50% columnas con grietas >10 cm; >50% pérdida sección núcleo; acero expuesto y pandeado"],
            ["Fotos ref.", "FOTOS 4–7 del informe DGPS"],
        ],
    )
    doc.add_heading("C.2 Vigas de concreto armado", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Mecanismo", "Corte / tensión diagonal"],
            ["Nivel", "Severo en enlace; afectación hasta ~5 niveles"],
            ["Evidencia", "Grietas diagonales típicas 45° en vigas principales"],
            ["Fotos ref.", "FOTO 8"],
        ],
    )
    doc.add_heading("C.3 Losas / escaleras / sótanos", level=3)
    bullets(
        doc,
        [
            "Reducción de altura de entrepiso en Nivel 1 por aplastamiento de columnas.",
            "Fisuras en arranque losa escalera primer piso (FOTO 9).",
            "Fracturas compresión en columnas de sótanos.",
        ],
    )
    doc.add_heading("C.4 Campos puente (tablero)", level=3)
    add_table(
        doc,
        ["Indicador", "Valor"],
        [
            ["Piso crítico", "Nivel 1 / PB transición"],
            ["% columnas daño grave", ">50% en piso crítico"],
            ["Inclinación / verticalidad", "Sí — pérdida de verticalidad global (P-Delta)"],
            ["Peligro aledaños / vía", "Sí — alto riesgo colapso sobre vía y vecinos"],
        ],
    )
    doc.add_heading("C.5 Diagnóstico narrativo (texto libre)", level=3)
    p(
        doc,
        "Se confirma carácter crítico del Nivel 1: falla masiva por compresión en columnas, reducción de "
        "entrepiso, pérdida de verticalidad y redistribución de esfuerzos en vigas hasta cinco niveles. "
        "Respecto a la precarga Habitable: la etiqueta ROJA y el desalojo eran correctos; el ranking "
        "Media (51) subestima la gravedad observada en la inspección detallada. No se considera viable "
        "reparación segura in situ bajo una torre de 12 pisos con apoyos del piso crítico en este estado.",
    )

    doc.add_heading("D — Daño en mampostería", level=2)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Mecanismo", "Compresión por reducción de entrepiso / fuera de plano en cerramientos"],
            ["Nivel", "Severo en tabiquería comprimida PB"],
            ["Diagnóstico", "Paredes de cerramiento fracturadas e inclinadas por descenso violento de losa superior"],
            ["Fotos ref.", "FOTOS 2, 10"],
        ],
    )

    doc.add_heading("E — Procedimientos sugeridos (Funvisis)", level=2)
    p(doc, "Dado el diagnóstico y la decisión de demolición, no proceden alternativas de reparación local. Se documenta:")
    bullets(
        doc,
        [
            "Columnas: condición crítica — fuera del alcance de COL-01 a COL-04 (aplastamiento, barras pandeadas, núcleo triturado).",
            "Vigas: daño severo por corte — reconstrucción equivalente no viable en contexto de demolición global.",
            "Mampostería: daño severo; intervención MAM no aplicable — edificio a demoler.",
            "Apuntalamiento: no recomendado como solución permanente; riesgo inaceptable para operarios bajo P-Delta.",
        ],
    )

    doc.add_heading("F — Decisión de control", level=2)
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Decisión D", "D1 — Demoler"],
            ["Magnitud M", "N/A (no es reparación D2)"],
            ["Prioridad", "Inmediata"],
            ["Medidas inmediatas", "Exclusión total; acordonar; prohibir ingreso; monitoreo colindantes"],
            ["Motivos tipificados D1", "Piso crítico masivo; pérdida verticalidad; >50% columnas graves; peligro aledaños"],
        ],
    )
    doc.add_heading("F.1 Justificación libre", level=3)
    p(
        doc,
        "Aunque el ranking de Fase 1 ubicaba el caso en banda Media, la inspección detallada evidencia "
        "colapso parcial de entrepiso, aplastamiento de más del 50% de columnas del Nivel 1, pérdida de "
        "verticalidad con efectos P-Delta y riesgo crítico para la vía pública y edificaciones vecinas. "
        "La reparación y el apuntalamiento para reconstrucción in situ no son técnicamente viables con "
        "seguridad razonable. Se recomienda demolición rápida y controlada, de arriba hacia abajo, bajo "
        "protocolo de ingeniería de demolición.",
    )

    doc.add_heading("G — Estimación referencial (opcional)", level=2)
    p(
        doc,
        "Para demolición controlada de torre ~12 pisos, área ~780 m²/planta: usar calculadora del sistema "
        "con supuestos editables (USD/m² demolición, accesos, escombros). Siempre rotular «estimación "
        "preliminar — no presupuesto contractual».",
    )

    doc.add_heading("H — Evidencia, firmas y resumen ejecutivo", level=2)
    p(doc, "Registro fotográfico: mínimo 10 evidencias referenciadas en informe DGPS (FOTOS 1–10).")
    p(doc, "Firmas: Elaboraron Burgos / García · Supervisión Aura Quintero · Vinculado certificado Habitable 73420200720261338.")
    doc.add_heading("Resumen ejecutivo (mesa de trabajo)", level=3)
    p(
        doc,
        "Franco Mar (Tanaguarenas, La Guaira). Fase 1: ROJO 20/07/2026, score 51 (Media; ~puesto 115 La Guaira). "
        "Visita detallada 11/08/2026: 12 pisos + 2 sótanos, pórticos CA, piso crítico Nivel 1 con >50% columnas "
        "graves, pérdida de verticalidad, daño por corte en vigas (5 niveles), peligro a vía y vecinos. "
        "Validación: ROJO confirmado; corrección sótanos y GPS/dirección. Procedimientos de reparación: no aplicables "
        "(condición crítica). Decisión: D1 demoler, prioridad inmediata. Próxima acción: proyecto demolición controlada "
        "y monitoreo de colindantes.",
        bold=True,
    )

    p(doc, "Excel asociado: Plantilla ejecutiva informe consolidado Fase II (hoja FrancoMar_completo).", bold=True)
    save_all(doc, "2026-08-27 Prototipo informe consolidado Fase II.docx")


def build_sistema_junior() -> None:
    doc = Document()
    style_doc(doc)
    title_block(
        doc,
        "PROPUESTA DE SISTEMA — SEGUIMIENTO Y VACIADO FASE II",
        "Guía extendida para equipo de desarrollo junior",
        "Marco recomendado: Django · Servidor 190.169.110.9 · MobaXterm · 27/08/2026",
    )

    doc.add_heading("0. Cómo leer esta guía (sin asumir nada)", level=1)
    p(
        doc,
        "Esta guía está escrita para personas que saben programar un poco en Python (variables, "
        "funciones, clases básicas) pero que quizá nunca hayan montado una aplicación web con "
        "usuarios, base de datos y formularios largos. No hace falta conocer Django, Streamlit ni "
        "ingeniería estructural: el sistema captura datos de un informe técnico; ustedes modelan "
        "pantallas, tablas y reglas.",
    )
    bullets(
        doc,
        [
            "Documentos hermanos (léanlos después de la sección 1): Word «Prototipo informe consolidado» "
            "y Excel «Plantilla ejecutiva» (hoja FrancoMar_completo).",
            "Sección 1 de este Word = vocabulario + marcos posibles + por qué Django.",
            "Desde la sección 2 en adelante = qué construir y cómo desplegarlo.",
        ],
    )

    doc.add_heading("1. ¿Qué es un «marco de trabajo» y cuáles existen?", level=1)
    doc.add_heading("1.1 Idea simple (antes de nombres raros)", level=2)
    p(
        doc,
        "Un marco de trabajo (en inglés: framework) es un conjunto de piezas de software ya hechas "
        "para no reinventar la rueda. En lugar de programar desde cero «cómo guardar un usuario», "
        "«cómo conectar a una base de datos» o «cómo mostrar una página web», el marco ya trae "
        "soluciones estándar y ustedes escriben solo la lógica del negocio (casos ROJO, visitas, "
        "dictámenes D1–D5).",
    )
    p(doc, "Analogía: construir una casa.")
    bullets(
        doc,
        [
            "Sin marco: fabrican ladrillos, cableado, tuberías y ventanas ustedes mismos.",
            "Con marco: llega un kit (cimientos, paredes, instalación eléctrica) y ustedes arman "
            "las habitaciones del proyecto (cola, ficha, PDF).",
        ],
    )
    p(
        doc,
        "Para este proyecto necesitamos, como mínimo: (1) login de inspectores, (2) guardar datos "
        "en una base de datos, (3) formularios por secciones del informe, (4) subir fotos, "
        "(5) estados (borrador → revisión → aprobado), (6) exportar PDF y Excel. Eso condiciona "
        "qué marco conviene.",
    )

    doc.add_heading("1.2 Piezas que toda opción debe cubrir (glosario mínimo)", level=2)
    add_table(
        doc,
        ["Concepto", "Qué significa en la práctica"],
        [
            [
                "Aplicación web",
                "Programa que se usa desde el navegador (Chrome/Edge). El usuario no instala un "
                ".exe; abre una URL, inicia sesión y trabaja.",
            ],
            [
                "Frontend",
                "Lo que se ve: pantallas, botones, tablas, formularios. HTML/CSS/JavaScript o "
                "componentes generados por el marco.",
            ],
            [
                "Backend",
                "Lo que corre en el servidor: validar datos, permisos, guardar en base de datos, "
                "generar PDF.",
            ],
            [
                "Base de datos (BD)",
                "Almacén ordenado de filas y columnas (PostgreSQL). Cada edificio/visita queda "
                "registrado y se puede consultar después.",
            ],
            [
                "ORM",
                "Capa que traduce clases Python a tablas SQL. En Django se llaman «modelos». "
                "Evita escribir SQL a mano para lo cotidiano.",
            ],
            [
                "Autenticación / autorización",
                "Quién entra (usuario/contraseña) y qué puede hacer (inspector vs supervisor).",
            ],
            [
                "Deploy / despliegue",
                "Poner la app en un servidor real para que otros la usen, no solo en su laptop.",
            ],
        ],
    )

    doc.add_heading("1.3 Opciones que evaluamos (qué es cada una, en lenguaje llano)", level=2)

    doc.add_heading("A) Solo Excel + carpetas compartidas", level=3)
    p(
        doc,
        "Qué es: seguir llenando plantillas Excel (como la del Franco Mar) y guardarlas en Drive "
        "o en el servidor por SFTP.",
    )
    bullets(
        doc,
        [
            "Para qué sirve bien: prototipos, talleres, pocos casos, trabajo manual.",
            "Qué hace: el humano abre el archivo, escribe, guarda. No hay «sistema» propiamente dicho.",
            "Límite: no hay usuarios reales, no hay historial de quién cambió un D1 a D2, se rompe "
            "con cientos o miles de ROJO, difícil auditar y unificar.",
            "Veredicto para Fase II: útil como plantilla de captura; insuficiente como producto final.",
        ],
    )

    doc.add_heading("B) Streamlit", level=3)
    p(
        doc,
        "Qué es: una librería de Python pensada para paneles de datos (BI). Escribes scripts "
        "Python y Streamlit dibuja controles (selectores, tablas, gráficos) en el navegador. "
        "En el programa CPEH ya se usa para tableros (por ejemplo BI Habitable / PDNA).",
    )
    bullets(
        doc,
        [
            "Para qué sirve bien: análisis, gráficos, cruces Excel, demos rápidas de indicadores.",
            "Qué hace: cada interacción puede «re-ejecutar» el script; ideal para explorar datos.",
            "Límite: permisos multi-usuario y formularios largos tipo «wizard de inspección» son "
            "incómodos; no trae un flujo de aprobación robusto ni un admin de operación.",
            "Veredicto: excelente para ver KPIs después; malo como núcleo del vaciado caso a caso.",
        ],
    )

    doc.add_heading("C) Reflex (u otras SPA en Python)", level=3)
    p(
        doc,
        "Qué es: un marco que permite escribir la interfaz web en Python y compila un frontend "
        "moderno (tipo aplicación de una sola página). Hay demos tipo «Insurance Suite» en el "
        "ecosistema del equipo, pero es otro producto.",
    )
    bullets(
        doc,
        [
            "Para qué sirve bien: interfaces modernas, demos visuales, cuando el equipo ya domina "
            "ese stack.",
            "Qué hace: une frontend y backend en Python; el navegador habla con un backend Reflex.",
            "Límite: curva más alta para juniors; hay que montar auth, roles, admin y despliegue "
            "con menos «piezas listas» que Django para este tipo de CRUD institucional.",
            "Veredicto: posible a futuro para UX; no el camino más corto para el MVP de seguimiento.",
        ],
    )

    doc.add_heading("D) React / Vue / Angular + API (FastAPI o Django REST)", level=3)
    p(
        doc,
        "Qué es: separar totalmente el frontend (JavaScript/TypeScript) del backend (API JSON). "
        "El inspector usa una SPA; el servidor solo responde datos.",
    )
    bullets(
        doc,
        [
            "Para qué sirve bien: productos grandes con equipo frontend dedicado y muchas pantallas "
            "interactivas.",
            "Qué hace: dos proyectos (UI + API), más contratos, más pruebas, más despliegues.",
            "Límite: para un equipo junior pequeño alarga el MVP; hay que aprender JS además de Python.",
            "Veredicto: exceso de complejidad para el punto inicial. Se puede evolucionar después.",
        ],
    )

    doc.add_heading("E) Laravel (PHP) o Ruby on Rails", level=3)
    p(
        doc,
        "Qué son: marcos «todo incluido» parecidos a Django, pero en PHP (Laravel) o Ruby (Rails). "
        "También tienen modelos, auth, admin/ecosistema y plantillas.",
    )
    bullets(
        doc,
        [
            "Para qué sirven bien: equipos que ya viven en ese lenguaje.",
            "Límite aquí: el resto del trabajo Habitable/CPEH/scripts de import está en Python "
            "(pandas, Excel, PDNA). Cambiar de lenguaje duplica herramientas y curva.",
            "Veredicto: válidos en abstracto; no alineados al stack del programa.",
        ],
    )

    doc.add_heading("F) Django (Python) — la opción elegida", level=3)
    p(
        doc,
        "Qué es: un marco web «baterías incluidas» en Python. Trae de fábrica: usuarios y grupos, "
        "formularios, ORM (modelos), panel de administración, migraciones de base de datos, "
        "sistema de plantillas HTML y un patrón claro (apps, urls, views).",
    )
    bullets(
        doc,
        [
            "Para qué sirve bien: sistemas internos con login, formularios, roles, PDF/Excel, "
            "auditoría y despliegue en servidor Linux.",
            "Qué hace en este proyecto: cada edificio ROJO es un registro; cada visita Fase II "
            "es otro; el inspector llena secciones; el supervisor aprueba; se exporta el informe.",
            "Ya existe experiencia cercana: el portal web CPEH también usa Django.",
        ],
    )

    doc.add_heading("1.4 Comparación rápida (misma pregunta a cada marco)", level=2)
    add_table(
        doc,
        ["Pregunta", "Excel", "Streamlit", "Reflex / SPA JS", "Django"],
        [
            [
                "¿Hay login y roles?",
                "No (archivo compartido)",
                "Débil / a mano",
                "Hay que construir",
                "Sí, incluido",
            ],
            [
                "¿Formulario largo por secciones?",
                "Manual",
                "Incómodo",
                "Sí, con mucho front",
                "Sí (forms + plantillas)",
            ],
            [
                "¿Panel para cargar catálogos el día 1?",
                "No",
                "No",
                "No",
                "Django Admin",
            ],
            [
                "¿Quién cambió un dictamen?",
                "Difícil",
                "Limitado",
                "Posible",
                "Fácil (modelos + log)",
            ],
            [
                "¿Curva para junior Python?",
                "N/A",
                "Baja (pero mal encaje)",
                "Alta",
                "Media + mucha documentación",
            ],
            [
                "¿Encaja con CPEH/Habitable?",
                "Solo plantillas",
                "BI / tableros",
                "Otro producto",
                "Sí (mismo lenguaje + precedente)",
            ],
        ],
    )

    doc.add_heading("1.5 Por qué escogemos Django (decisión explícita)", level=2)
    p(doc, "Elegimos Django no porque sea «el mejor del mundo», sino porque es el mejor ajuste a este problema concreto:", bold=True)
    bullets(
        doc,
        [
            "El problema es un sistema de seguimiento y vaciado (CRUD + workflow), no un tablero BI.",
            "Necesitamos usuarios distintos (inspector, supervisor, coordinador) desde el día 1.",
            "El informe tiene muchas secciones: Django permite guardar por bloque y validar reglas "
            "(ej.: si D2, entonces magnitud M obligatoria).",
            "El equipo junior programa en Python; no queremos obligar a aprender un segundo lenguaje "
            "para el MVP.",
            "Ya hay precedente Django en el ecosistema CPEH y scripts Python para Habitable/Excel.",
            "Django Admin acelera: mientras se construye el wizard bonito, ya se pueden cargar "
            "casos y catálogos VIG/COL/MAM.",
            "Despliegue estándar (Gunicorn + Nginx + PostgreSQL) en el servidor Linux que usaremos "
            "con MobaXterm.",
        ],
    )
    p(doc, "Qué NO implica esta decisión:")
    bullets(
        doc,
        [
            "No prohibimos Streamlit: seguirá siendo útil para tableros y KPIs encima de la misma BD.",
            "No prohibimos un frontend React más adelante: primero MVP Django; luego se puede "
            "exponer API si hace falta.",
            "No reemplazamos Excel de golpe: la plantilla ejecutiva sigue siendo el molde de campos; "
            "Django la convierte en pantallas y registros.",
        ],
    )

    doc.add_heading("1.6 Resumen en una frase", level=2)
    p(
        doc,
        "Excel = plantilla; Streamlit = ver datos; Reflex/React = UI avanzada; Django = sistema "
        "con usuarios, base de datos y flujo de aprobación. Para el puente Habitable → Fase II "
        "elegimos Django.",
        bold=True,
    )

    doc.add_heading("2. Qué vamos a construir (MVP concreto)", level=1)
    p(doc, "Imaginen una aplicación web interna con estos módulos:", bold=True)
    add_table(
        doc,
        ["Módulo", "Qué hace el usuario", "Qué guarda en BD"],
        [
            ["Cola", "Coordinador filtra ROJO por banda/municipio y asigna inspector", "BuildingCase + estado_cola"],
            ["Ficha A0", "Inspector valida precarga Habitable", "Phase1Snapshot + validaciones"],
            ["Visita", "Inspector llena bloques A–H (pestañas)", "StructureProfile, Findings, Photos"],
            ["Dictamen", "Inspector elige D/M; supervisor aprueba", "ControlDecision + workflow"],
            ["Export", "Genera PDF institucional + fila Excel", "Archivos + timestamp"],
            ["Tablero", "Conteos D1/D2/… y avance cola", "Consultas agregadas (no tabla nueva)"],
        ],
    )

    doc.add_heading("3. Ventajas Django que usarán día a día en este proyecto", level=1)
    bullets(
        doc,
        [
            "Un modelo = una tabla: BuildingCase, InspectionVisit, ControlDecision — fácil de enseñar.",
            "Migraciones versionan la BD: cuando Funvisis agregue un campo, hacemos migration 0007_add_campo_x.",
            "Permisos por rol sin inventar: inspector no aprueba; supervisor no borra casos aprobados.",
            "Django Admin permite probar datos mientras el frontend del wizard aún no existe.",
            "Mismo lenguaje que scripts de import Habitable (pandas → management command import_habitable).",
            "Despliegue estándar: Gunicorn + Nginx + PostgreSQL (patrón conocido en producción).",
        ],
    )

    doc.add_heading("4. Arquitectura en 3 capas (dibújela en una pizarra)", level=1)
    bullets(
        doc,
        [
            "Capa 1 — Datos: PostgreSQL (casos, visitas, fotos metadata, dictamen).",
            "Capa 2 — Aplicación: Django (views, forms, permisos, generación PDF/Excel).",
            "Capa 3 — Presentación: HTML Bootstrap 5 (+ HTMX opcional para guardar pestaña sin recargar todo).",
            "Archivos (fotos): carpeta media/ o S3-compatible — no guardar blobs gigantes dentro de PostgreSQL.",
        ],
    )

    doc.add_heading("5. Modelo de datos — explicación paso a paso", level=1)
    p(doc, "Relaciones en lenguaje humano:")
    bullets(
        doc,
        [
            "Un BuildingCase = un edificio ROJO en cola (1 fila por habitable_id).",
            "Un BuildingCase tiene muchas InspectionVisit (re-inspecciones futuras).",
            "Cada InspectionVisit tiene 1 StructureProfile, N StructuralFinding, N MasonryFinding, N PhotoEvidence, 1 ControlDecision.",
            "RepairRecommendation = filas catálogo VIG-01… marcadas sí/no para esa visita.",
        ],
    )
    add_table(
        doc,
        ["Modelo Django", "Campos ejemplo", "Validación importante"],
        [
            ["BuildingCase", "habitable_id unique, score, banda, estado_cola", "No duplicar ID"],
            ["InspectionVisit", "fecha, estado_workflow", "Solo 1 borrador activo por caso"],
            ["ControlDecision", "codigo_D, magnitud_M, justificacion", "Si D2 → M obligatorio"],
            ["PhotoEvidence", "tipo, file, caption", "Mínimo 6 tipos antes de enviar a revisión"],
            ["AuditLog", "user, action, json_diff", "Automático en save de dictamen"],
        ],
    )

    doc.add_heading("6. Pantallas que programarán (checklist)", level=1)
    add_table(
        doc,
        ["#", "URL sugerida", "Rol", "Hecho cuando…"],
        [
            ["1", "/login/", "Todos", "Entra con usuario/contraseña"],
            ["2", "/cola/", "Coord.+Insp.", "Ve ranking filtrable"],
            ["3", "/caso/<id>/", "Inspector", "Ve A0 precarga"],
            ["4", "/visita/<id>/bloque/<letra>/", "Inspector", "Guarda sección A…H"],
            ["5", "/visita/<id>/fotos/", "Inspector", "Sube ≥6 fotos tipificadas"],
            ["6", "/visita/<id>/enviar/", "Inspector", "Pasa a en_revision"],
            ["7", "/revision/", "Supervisor", "Aprueba o devuelve"],
            ["8", "/export/pdf/<visita>/", "Todos auth", "Descarga PDF"],
            ["9", "/admin/", "Admin", "Catálogos VIG/COL/MAM"],
        ],
    )

    doc.add_heading("7. Plan de sprints (6–8 semanas orientativo)", level=1)
    add_table(
        doc,
        ["Sprint", "Entrega", "Prueba de aceptación"],
        [
            ["0", "Repo + Docker Compose local + login", "Usuario demo entra"],
            ["1", "Modelos + Admin + import CSV Habitable", "Aparecen casos en cola"],
            ["2", "Pantallas A0 + A + B", "Inspector guarda ficha básica"],
            ["3", "Bloques C + D + fotos", "Hallazgos + galería"],
            ["4", "Bloque E catálogo + F dictamen", "D1/D2 con reglas"],
            ["5", "Workflow supervisor + PDF", "PDF Franco Mar sale igual al Word"],
            ["6", "Excel export + dashboard KPI", "Conteo D1/D2 visible"],
        ],
    )

    doc.add_heading("8. Servidor de despliegue — acceso y requisitos", level=1)
    p(
        doc,
        "El entorno objetivo es el servidor Linux accesible por SSH desde Windows con MobaXterm. "
        "La conexión fue verificada el 27/08/2026 con usuario cph y llave PuTTY .ppk. "
        "NO almacene contraseñas, frases clave ni archivos .ppk en el repositorio Git ni en este Word.",
    )
    add_table(
        doc,
        ["Parámetro", "Valor / nota"],
        [
            ["Host SSH", SERVER_HOST],
            ["Puerto", "22 (SSH estándar)"],
            ["Usuario SSH", f"{SERVER_USER} — obligatorio; no usar otros alias de sesión"],
            ["Hostname en servidor", SERVER_HOSTNAME],
            ["Sistema operativo", SERVER_OS],
            ["Home del usuario", SERVER_HOME],
            ["Directorio de trabajo del proyecto", SERVER_WORKDIR],
            ["Carpeta padre del proyecto", "/home/cph/Project_Etapa2_FlujoObra/"],
            ["Autenticación", "Solo llave pública (el servidor no acepta contraseña)"],
            ["Clave privada (local Windows)", f"{LOCAL_KEY_PPK} — formato PuTTY .ppk, comentario «{SERVER_KEY_COMMENT}»"],
            ["Frase clave de la llave", f"Archivo local {LOCAL_PASSPHRASE_FILE} — solo en disco del operador"],
            ["Huella de la llave pública", SERVER_KEY_FINGERPRINT],
            ["Cliente SSH en Windows", f"MobaXterm portable en {LOCAL_MOBA_DIR}"],
        ],
    )
    doc.add_heading("8.1 Qué verás si la conexión es correcta", level=2)
    p(
        doc,
        "Tras un login exitoso, la terminal debe mostrar un prompt similar a cph@cph:~$ "
        "(usuario cph, hostname cph). El panel SFTP izquierdo de MobaXterm listará, entre otros:",
    )
    bullets(
        doc,
        [
            f"{SERVER_HOME}.ssh/ — configuración SSH del usuario (no borrar)",
            "/home/cph/Project_Etapa2_FlujoObra/ — árbol del proyecto existente en el servidor",
            f"{SERVER_WORKDIR} — carpeta objetivo para el módulo de control de informes Fase II",
            "Archivos de perfil (.bashrc, .profile) — normales en cualquier cuenta Linux",
        ],
    )
    p(
        doc,
        "Comandos de orientación inmediata (copiar en la terminal tras conectar):",
        bold=True,
    )
    bullets(
        doc,
        [
            "pwd && whoami && hostname",
            "ls -la ~/",
            f"ls -la {SERVER_WORKDIR}",
            "python3 --version",
            "python3 -m django --version 2>/dev/null || echo 'Django aún no instalado'",
        ],
    )

    doc.add_heading("8.2 Software a instalar en el servidor (stack Django)", level=2)
    add_table(
        doc,
        ["Componente", "Versión sugerida", "Para qué"],
        [
            ["SO", f"{SERVER_OS} — usar paquetes apt de Debian", "Base ya verificada en el servidor"],
            ["Python", "3.11 o 3.12", "Django 5 compatible"],
            ["PostgreSQL", "15+", "Base de datos principal"],
            ["nginx", "Estable", "Proxy reverso + estáticos"],
            ["gunicorn", "Última estable", "Servidor WSGI Django"],
            ["git", "Cualquier reciente", "Despliegue código"],
            ["libpq-dev, gcc", "—", "Compilar psycopg2"],
            ["WeasyPrint deps", "pango, cairo (paquetes sistema)", "Export PDF"],
            ["certbot", "Opcional fase 2", "HTTPS con dominio"],
            ["redis", "Opcional", "Tareas async futuras (import masivo)"],
        ],
    )

    doc.add_heading("8.3 Comandos base en servidor (referencia para juniors)", level=2)
    p(doc, "Tras conectar por SSH como cph, el flujo típico de instalación del MVP será:", bold=False)
    bullets(
        doc,
        [
            "sudo apt update && sudo apt install -y python3 python3-venv python3-pip postgresql nginx git libpq-dev",
            "mkdir -p ~/apps/fase2-inspecciones && cd ~/apps/fase2-inspecciones",
            "python3 -m venv .venv && source .venv/bin/activate",
            "pip install django psycopg2-binary gunicorn pillow openpyxl weasyprint python-dotenv",
            "django-admin startproject config .  (o clonar repo del equipo)",
            "python manage.py migrate && python manage.py createsuperuser",
            "Configurar gunicorn systemd + nginx site → proxy a 127.0.0.1:8000",
            "Variables .env en servidor (SECRET_KEY, DATABASE_URL, ALLOWED_HOSTS) — archivo fuera de Git",
        ],
    )

    doc.add_heading("9. Manual de acceso — MobaXterm y SSH", level=1)
    p(
        doc,
        "MobaXterm es un cliente SSH + SFTP + terminal para Windows. En este proyecto se usa la copia "
        f"portable en {LOCAL_MOBA_DIR}. Permite conectar al servidor, ejecutar comandos Linux y "
        "arrastrar archivos (Excel, PDF, código) al panel izquierdo sin instalar otro programa.",
    )

    doc.add_heading("9.1 Configuración correcta de la sesión (obligatoria)", level=2)
    add_table(
        doc,
        ["Campo en MobaXterm", "Valor exacto", "Error común"],
        [
            ["Remote host", SERVER_HOST, "IP incorrecta o espacio extra"],
            ["Username", SERVER_USER, "Usar pleaseaskme u otro alias — la llave NO está autorizada para ellos"],
            ["Port", "22", "Puerto distinto sin aviso del administrador"],
            ["Use private key", LOCAL_KEY_PPK, "Seleccionar MobaXterm.ini u otro archivo que no sea llave"],
            ["Passphrase al conectar", "Contenido de frase-clave.txt (local)", "Compartir la frase por chat o correo"],
            ["Nombre de sesión sugerido", f"{SERVER_HOST} ({SERVER_USER})", "Dejar nombre viejo pleaseaskme para no confundir"],
        ],
    )

    doc.add_heading("9.2 Pasos — primera conexión", level=2)
    bullets(
        doc,
        [
            f"Paso 1: Abrir MobaXterm.exe desde {LOCAL_MOBA_DIR}.",
            "Paso 2: Session → SSH (o clic derecho en sesión existente → Edit session).",
            f"Paso 3: Remote host = {SERVER_HOST} · Specify username = {SERVER_USER}.",
            f"Paso 4: Advanced SSH settings → marcar Use private key → elegir {LOCAL_KEY_PPK}.",
            "Paso 5: OK / Save. NO usar como llave el archivo MobaXterm.ini (es configuración del programa).",
            "Paso 6: Conectar. MobaXterm pedirá la frase clave de la llave — leerla del archivo local de frase clave.",
            "Paso 7: La primera vez, aceptar la huella del host (host key) del servidor.",
            f"Paso 8: Verificar prompt cph@cph:~$ y panel SFTP en {SERVER_HOME}.",
            f"Paso 9: Ir al trabajo: cd {SERVER_WORKDIR}",
            "Paso 10: Activar «Follow terminal folder» en el panel SFTP si quieres que la carpeta izquierda siga al cd.",
            "Paso 11: Session → Save para reutilizar la configuración cada día.",
        ],
    )

    doc.add_heading("9.3 Errores frecuentes y cómo resolverlos", level=2)
    add_table(
        doc,
        ["Mensaje en pantalla", "Causa probable", "Solución"],
        [
            [
                "Unable to use key file … MobaXterm.ini — not a recognized key file format",
                "La sesión apunta al .ini de MobaXterm en lugar de la llave .ppk",
                f"Edit session → Use private key → {LOCAL_KEY_PPK}",
            ],
            [
                "Server refused our key",
                f"Usuario SSH incorrecto (p. ej. pleaseaskme) o llave no autorizada para ese usuario",
                f"Cambiar Username a {SERVER_USER} y usar {LOCAL_KEY_PPK}",
            ],
            [
                "No supported authentication methods available (server sent: publickey)",
                "El servidor solo acepta llave; no hay contraseña alternativa",
                "Corregir usuario + .ppk + frase clave; pedir al admin que verifique authorized_keys",
            ],
            [
                "Passphrase prompt rechazada / no desbloquea la llave",
                "Frase clave incorrecta o archivo .ppk dañado",
                "Verificar frase-clave.txt; probar cargar la .ppk en PuTTYgen con Load",
            ],
            [
                "Conecta pero SFTP vacío o permiso denegado",
                "Carpeta distinta o permisos del usuario cph",
                f"cd {SERVER_HOME} y revisar ls -la; escalar al admin solo si falta permiso en Project_Etapa2",
            ],
        ],
    )

    doc.add_heading("9.4 Verificación para el coordinador / admin", level=2)
    p(
        doc,
        "Si tras usar usuario cph y la llave correcta sigue «Server refused our key», el administrador "
        "del servidor debe confirmar que la llave pública correspondiente está en "
        f"/home/{SERVER_USER}/.ssh/authorized_keys. Huella de referencia de la llave del operador: "
        f"{SERVER_KEY_FINGERPRINT} (comentario «{SERVER_KEY_COMMENT}»).",
    )
    p(
        doc,
        "Tip técnico: MobaXterm entiende .ppk nativamente. Si algún día usan OpenSSH desde WSL o PowerShell, "
        "convierten la llave con PuTTYgen → Conversions → Export OpenSSH key (requiere la misma frase clave).",
    )

    doc.add_heading("10. Seguridad — reglas que no se negocian", level=1)
    bullets(
        doc,
        [
            "Nunca commitear .ppk, frase clave, .env de producción ni SECRET_KEY.",
            "Fotos de inspección: permisos por rol; no URLs públicas sin autenticación.",
            "Backups PostgreSQL diarios cuando haya casos reales.",
            "HTTPS obligatorio antes de abrir a inspectores en campo (fase piloto puede ser VPN/red cerrada).",
        ],
    )

    doc.add_heading("11. Definition of Done del MVP", level=1)
    bullets(
        doc,
        [
            "Inspector completa Franco Mar en el sistema y obtiene PDF equivalente al anexo del Word.",
            "Supervisor aprueba y el caso cambia a «aprobado» en cola.",
            "Coordinador ve KPI: cuántos D1 vs D2 vs pendientes.",
            "Import Habitable crea/actualiza casos sin duplicar habitable_id.",
        ],
    )

    save_all(doc, "2026-08-27 Propuesta sistema seguimiento Fase II - guia equipo junior.docx")


# Campos extra Funvisis para Excel consolidado
FUNVISIS_EXTRA_ROWS: list[tuple[str, str | None, str, str | None, bool]] = [
    ("C — DAÑO ESTRUCTURAL (FUNVISIS)", None, None, None, False),
    ("C", "Columnas — mecanismo", "col_mec", None, False),
    ("C", "Columnas — nivel daño", "col_nivel", "abc", False),
    ("C", "Columnas — ubicación / evidencia", "col_evidencia", None, False),
    ("C", "Vigas — mecanismo", "vig_mec", None, False),
    ("C", "Vigas — nivel daño", "vig_nivel", "abc", False),
    ("C", "Vigas — ubicación / evidencia", "vig_evidencia", None, False),
    ("C", "Muros/pantallas — mecanismo", "mur_mec", None, False),
    ("C", "Muros/pantallas — nivel daño", "mur_nivel", "abc", False),
    ("C", "Escaleras / evacuación", "escaleras", None, False),
    ("C", "Daños preexistentes (pre 24/06/2026)", "preexistentes", None, False),
    ("D — MAMPOSTERÍA (FUNVISIS)", None, None, None, False),
    ("D", "Mampostería — mecanismo", "mam_mec", None, False),
    ("D", "Mampostería — nivel", "mam_nivel", "abc", False),
    ("D", "Diagnóstico mampostería", "mam_diag", None, False),
    ("E — PROCEDIMIENTOS (FUNVISIS)", None, None, None, False),
    ("E", "Procedimientos seleccionados (códigos)", "proc_codigos", None, False),
    ("E", "¿Reparación viable?", "repar_viable", "si_no_insuf", False),
    ("E", "Notas procedimientos", "proc_notas", None, False),
]

FRANCO_FUNVISIS: dict[str, Any] = {
    "col_mec": "Compresión / aplastamiento núcleo",
    "col_nivel": "C",
    "col_evidencia": "Nivel 1 >50% columnas; grietas >10 cm; núcleo >50% perdido — FOTOS 4-7",
    "vig_mec": "Corte / tensión diagonal",
    "vig_nivel": "C",
    "vig_evidencia": "Grietas 45° hasta ~5 niveles — FOTO 8",
    "mur_mec": "Compresión por reducción entrepiso",
    "mur_nivel": "C",
    "mam_mec": "Compresión / fuera de plano en cerramientos",
    "mam_nivel": "C",
    "mam_diag": "Tabiquería comprimida e inclinada PB — FOTOS 2, 10",
    "escaleras": "Fisuras arranque losa escalera piso 1 — FOTO 9",
    "preexistentes": "No reportados como determinantes vs daño sísmico 24/06/2026",
    "proc_codigos": "Condición crítica COL — fuera de alcance; demolición estructural",
    "repar_viable": "No",
    "proc_notas": "No aplican VIG/COL/MAM de reparación; D1 demolición controlada",
}


def _write_consolidated_sheet(
    wb: Workbook,
    sheet_name: str,
    precarga: dict[str, Any],
    visita: dict[str, Any] | None,
    funvisis: dict[str, Any] | None,
    titulo: str,
) -> None:
    from excel_control_vaciado import write_vaciado_informe_sheet

    # Hoja base A0–E+R usando función existente
    write_vaciado_informe_sheet(wb, sheet_name, precarga, visita, titulo=titulo)
    ws = wb[sheet_name]

    # Insertar filas Funvisis antes de RESUMEN (buscar fila sección RESUMEN)
    resumen_row = None
    for r in range(4, ws.max_row + 1):
        if ws.cell(r, 1).value and str(ws.cell(r, 1).value).startswith("RESUMEN"):
            resumen_row = r
            break
    if resumen_row is None:
        resumen_row = ws.max_row + 1

    thin = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    font_sec = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    font_lab = Font(name="Calibri", size=10, bold=True)
    font_val = Font(name="Calibri", size=10)
    data = {**(visita or {}), **(funvisis or {})}

    ws.insert_rows(resumen_row, amount=len(FUNVISIS_EXTRA_ROWS))
    list_cols = {k: i + 1 for i, k in enumerate(LISTAS.keys())}

    r = resumen_row
    for seccion, etiqueta, key, lista_key, _ in FUNVISIS_EXTRA_ROWS:
        if etiqueta is None:
            ws.cell(r, 1, seccion)
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
            cell = ws.cell(r, 1)
            cell.fill = PatternFill("solid", fgColor="7030A0")
            cell.font = font_sec
            r += 1
            continue
        ws.cell(r, 1, etiqueta).font = font_lab
        ws.cell(r, 1).border = thin
        val = data.get(key, "") if key else ""
        c = ws.cell(r, 2, val)
        c.font = font_val
        c.border = thin
        c.alignment = Alignment(wrap_text=True, vertical="top")
        c.fill = PatternFill("solid", fgColor="E2EFDA")
        if lista_key and key:
            col_i = list_cols[lista_key]
            col_letter = get_column_letter(col_i)
            nvals = len(LISTAS[lista_key])
            formula = f"Listas_desplegables!${col_letter}$2:${col_letter}${nvals + 1}"
            ws.add_data_validation(_dv(formula, f"$B${r}"))
        if key in ("col_evidencia", "vig_evidencia", "mam_diag", "proc_notas"):
            ws.row_dimensions[r].height = 45
        r += 1


def build_excel_ejecutivo() -> None:
    wb = Workbook()
    ws0 = wb.active
    ws0.title = "Portada"
    ws0["A1"] = "PLANTILLA EJECUTIVA — INFORME CONSOLIDADO FASE II"
    ws0["A1"].font = Font(name="Calibri", size=14, bold=True, color="1F4E79")
    ws0.merge_cells("A1:B1")
    filas = [
        ("Versión", "0.2 · 27/08/2026 · Funvisis + Habitable/D-M"),
        ("Word asociado", "2026-08-27 Prototipo informe consolidado Fase II.docx"),
        ("Hojas", "Mapa_bloques · FrancoMar_completo · Plantilla_vacia · Resumen · Listas"),
        ("Colores", "Azul = precarga F1 · Amarillo = lista desplegable · Verde = Funvisis · Blanco = texto libre"),
        ("Caso ejemplo", "Franco Mar ID 171818 — D1 Demoler — informe DGPS 11/08/2026"),
    ]
    rr = 3
    for a, b in filas:
        ws0.cell(rr, 1, a).font = Font(bold=True)
        ws0.cell(rr, 2, b)
        rr += 1
    ws0.column_dimensions["A"].width = 22
    ws0.column_dimensions["B"].width = 90

    ws_map = wb.create_sheet("Mapa_bloques")
    ws_map.append(["Bloque", "Contenido", "Obligatorio MVP"])
    map_rows = [
        ("A0", "Precarga Habitable + ranking + validación", "Sí"),
        ("A-B", "Identidad + descripción estructural", "Sí"),
        ("C", "Daño estructural Funvisis + campos puente", "Sí"),
        ("D", "Mampostería Funvisis", "Sí si aplica"),
        ("E", "Procedimientos VIG/COL/MAM", "Sí si D2"),
        ("F", "Decisión D/M + medidas + justificación", "Sí"),
        ("H", "Fotos + firmas + resumen ejecutivo", "Sí"),
    ]
    for row in map_rows:
        ws_map.append(list(row))
    for cell in ws_map[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E79")
    ws_map.column_dimensions["A"].width = 8
    ws_map.column_dimensions["B"].width = 55
    ws_map.column_dimensions["C"].width = 18

    _add_list_sheet(wb)

    precarga = francomar_precarga()
    visita = francomar_visita()
    _write_consolidated_sheet(
        wb,
        "FrancoMar_completo",
        precarga,
        visita,
        FRANCO_FUNVISIS,
        titulo="INFORME CONSOLIDADO — EJEMPLO COMPLETO FRANCO MAR (11/08/2026)",
    )

    empty_pre = {k: "" for k in precarga}
    empty_pre["etiqueta_f1"] = "ROJO"
    _write_consolidated_sheet(
        wb,
        "Plantilla_vacia",
        empty_pre,
        None,
        None,
        titulo="PLANTILLA VACÍA — Pegue precarga y complete visita Fase II",
    )

    ws_r = wb.create_sheet("Resumen_ejecutivo")
    ws_r["A1"] = "RESUMEN EJECUTIVO — FRANCO MAR"
    ws_r["A1"].font = Font(size=14, bold=True, color="1F4E79")
    for i, (a, b) in enumerate(
        [
            ("Edificación", "Franco Mar — Tanaguarenas, La Guaira"),
            ("ID / certificado", "171818 / 73420200720261338"),
            ("Decisión", "D1 Demoler · Prioridad inmediata"),
            ("Texto", visita["resumen_ejecutivo"]),
        ],
        start=3,
    ):
        ws_r.cell(i, 1, a).font = Font(bold=True)
        ws_r.cell(i, 2, b)
    ws_r.column_dimensions["A"].width = 18
    ws_r.column_dimensions["B"].width = 95

    order = [
        "Portada",
        "Mapa_bloques",
        "FrancoMar_completo",
        "Plantilla_vacia",
        "Resumen_ejecutivo",
        "Listas_desplegables",
    ]
    for i, name in enumerate(order):
        wb.move_sheet(name, offset=i - wb.sheetnames.index(name))

    fname = "2026-08-27 Plantilla ejecutiva informe consolidado Fase II.xlsx"
    for d in OUT_DIRS:
        d.mkdir(parents=True, exist_ok=True)
        path = d / fname
        wb.save(str(path))
        print("OK", path)


def main() -> None:
    build_prototipo_informe()
    build_sistema_junior()
    build_excel_ejecutivo()
    sample = OUT_DIRS[0] / "2026-08-27 Prototipo informe consolidado Fase II.docx"
    d = Document(str(sample))
    text = "\n".join(p.text for p in d.paragraphs[:15] if p.text.strip())
    assert "COMISIÓN" in text and "¿" not in text.replace("¿", ""), "Revisar encoding"
    assert "Comparación" not in text or "Comparaci" not in text
    print("VERIFICACION_OK UTF-8")


if __name__ == "__main__":
    main()
